"""Generates the DHL Document Fine-Tuning Reference DOCX — v2 (updated after dataset rebuild)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

DHL_RED   = RGBColor(0xD4, 0x05, 0x11)
NAVY      = RGBColor(0x1A, 0x1A, 0x2E)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY  = RGBColor(0x55, 0x55, 0x55)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DHL_RED
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = NAVY
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK_GREY
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10)
    return p

def code(text):
    p = doc.add_paragraph()
    p.style = "No Spacing"
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A1A2E")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    return t

def pagebreak():
    doc.add_page_break()

def sp():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("DHL Document Intelligence\nFine-Tuning Reference Guide")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = DHL_RED

sp()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run(
    "Document Classification + Splitting · Single VLM · Sequential Context\n"
    "Synthetic Dataset Design · Inference Pipeline · Post-Processing"
)
r2.font.size = Pt(12)
r2.font.color.rgb = MID_GREY

sp()
ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_p.add_run(
    f"Version 2.0  |  Generated: {datetime.date.today().strftime('%d %B %Y')}\n"
    "Dataset rebuilt with prev-page context, balanced START/CONTINUATION, per-PDF split"
).font.size = Pt(10)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1: PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")
body(
    "This project trains a single Vision Language Model (VLM) that processes DHL logistics "
    "document pages one at a time and outputs both the document class AND whether the page "
    "is the start of a new document or a continuation of the previous one. One model, one "
    "inference call per page, solves classification and splitting simultaneously."
)

sp()
h2("1.1 The Single Model Objective")
add_table(
    ["Task", "What the model outputs", "How it is used"],
    [
        ["Classification",
         "Document class\n(e.g. 'Commercial Invoice')",
         "Route document to correct downstream extraction pipeline"],
        ["Splitting",
         "START or CONTINUATION\nfor this page",
         "Detect document boundaries in a mixed multi-page PDF"],
        ["Both in one call",
         "'Commercial Invoice | START'\nor 'House Bill of Lading | CONTINUATION'",
         "Split the class portion for classification; split the position portion for boundary detection"],
    ]
)

sp()
h2("1.2 Real-World Input Scenarios")
add_table(
    ["Scenario", "Description", "Example sequence"],
    [
        ["Single 1-page doc",
         "One standalone document, one page",
         "CI → 'Commercial Invoice | START'"],
        ["Single 2-3 page doc",
         "One document that overflows to multiple pages",
         "p1 → 'CI | START'  p2 → 'CI | CONTINUATION'  p3 → 'CI | CONTINUATION'"],
        ["Combined batch PDF (5-10+ pages)",
         "Multiple complete documents concatenated. Rule: one document always fully completes before the next begins — pages never interleave.",
         "CI(p1) → 'CI|START'\nCI(p2) → 'CI|CONT'\nHBL(p3) → 'HBL|START'\nHBL(p4) → 'HBL|CONT'\nHBL*(p5) → 'HBL|START'  ← NEW HBL doc, same class!\nHBL*(p6) → 'HBL|CONT'\nCOO(p7) → 'COO|START'"],
    ]
)
sp()
body(
    "The hardest case is page 5 above: a new HBL document beginning immediately after "
    "a previous HBL's continuation page. Without sequence context the model cannot "
    "reliably detect this boundary. This is why every training example includes the "
    "previous page's classification as text context (see Section 4)."
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2: 12 DOCUMENT CLASSES
# ══════════════════════════════════════════════════════════════════════════
h1("2. The 12 Document Classes")

add_table(
    ["#", "Class Name", "Issued By", "Purpose", "Multi-page?"],
    [
        ["01", "Commercial Invoice",              "Shipper/Exporter",           "Customs valuation, duty assessment",                             "Yes (500 2-page PDFs)"],
        ["02", "House Bill of Lading",            "DHL/Freight Forwarder",      "Ocean freight contract with shipper",                            "Yes (300 2-page PDFs)"],
        ["03", "Certificate of Origin",           "Chamber of Commerce",        "Proves country of manufacture for tariff purposes",              "No — single page only"],
        ["04", "Shipper's Letter of Instruction", "Shipper/Exporter",           "Authorizes DHL to file EEI and prepare export docs",             "Yes (300 2-page PDFs)"],
        ["05", "Dangerous Goods Declaration",     "Shipper (DG certified)",     "Required for shipping hazmat by air/sea",                        "Yes (300 2-page PDFs)"],
        ["06", "Verified Gross Mass",             "Shipper",                    "SOLAS mandatory container weight declaration",                   "Yes (300 2-page PDFs)"],
        ["07", "House Airway Bill",               "DHL/Freight Forwarder",      "Air cargo contract with shipper",                                "No — single page only"],
        ["08", "Packing List",                    "Shipper/Exporter",           "Itemizes package contents for customs/receiver",                 "Yes (500 2-page PDFs)"],
        ["09", "Customs Declaration (CN23)",      "Shipper/Postal Service",     "Postal customs declaration",                                     "No — single page only"],
        ["10", "Cargo Manifest",                  "Carrier/Freight Forwarder",  "Lists all cargo on vessel/aircraft for customs",                 "Yes (500 2-page PDFs)"],
        ["11", "Import/Export License (EEI)",     "Government/CBP",             "Customs entry summary or export control license",                "Yes (300 2-page PDFs)"],
        ["12", "Power of Attorney",               "Shipper/Importer",           "Legal authorization for DHL to act as customs agent",           "No — single page only"],
    ]
)
sp()
body(
    "Note: Classes 03, 07, 09, 12 currently have no multi-page generator. "
    "They appear only as single-page documents in the dataset. "
    "Future work: generate 2-page variants to improve CONTINUATION representation for these classes."
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3: DATASET STRUCTURE
# ══════════════════════════════════════════════════════════════════════════
h1("3. Dataset Structure")

h2("3.1 Source Data (PDFs)")
add_table(
    ["Directory", "Content", "Count", "All START?"],
    [
        ["Synthetic_Data/<class>/pdfs/",          "Single-page PDFs, one document per file",                        "12,000 total\n(1,000/class)", "Yes"],
        ["Synthetic_Data_MultiPage/<class>/pdfs/", "Multi-page PDFs — same document overflows to 2 pages",          "3,300 total\n(300-500/class for 8 classes)", "p1=START\np2+=CONTINUATION"],
        ["Synthetic_Data_Splitting_v2/pdfs/",     "Combined packets — multiple complete docs merged into one PDF",  "5,000 packets\navg 5.6 pages/packet", "Mixed: every first page\nof each doc is START"],
    ]
)

sp()
h2("3.2 Training Data (JSONL) — Final Stats")
add_table(
    ["Split", "Examples", "START", "CONTINUATION", "Sources"],
    [
        ["train.jsonl", "29,744",  "21,409 (72%)", "8,335 (28%)",  "splitting_packet: 22,533\nmulti_doc: 4,801\nsingle_doc: 2,410"],
        ["val.jsonl",   " 3,723",  " 2,742 (74%)", "  981 (26%)",  "splitting_packet: 2,854\nmulti_doc: 570\nsingle_doc: 299"],
        ["test.jsonl",  " 3,626",  " 2,653 (73%)", "  973 (27%)",  "splitting_packet: 2,703\nmulti_doc: 632\nsingle_doc: 291"],
        ["TOTAL",       "37,093",  "26,804 (72%)", "10,289 (28%)", ""],
    ]
)

sp()
body(
    "Key design decisions: (1) single_doc capped at 250/class (3,000 total) — standalone pages "
    "are less realistic and were causing 90% START dominance. (2) Split is by individual PDF stem, "
    "not by class folder — ensures every class is represented in all three splits. "
    "(3) All 3 splits maintain the same 72-74%/26-28% START/CONTINUATION ratio."
)

sp()
h2("3.3 Class Distribution in Training Set")
add_table(
    ["Class", "Train Examples", "% of Train", "Note"],
    [
        ["Commercial Invoice",              "7,658", "25.7%", "500 multi-page + heavy template coverage"],
        ["Packing List",                    "6,386", "21.5%", "500 multi-page + heavy template coverage"],
        ["House Bill of Lading",            "3,183", "10.7%", "300 multi-page, now used in packets"],
        ["Dangerous Goods Declaration",     "1,964",  "6.6%", "300 multi-page, now used in packets"],
        ["House Airway Bill",               "1,955",  "6.6%", "Single-page only, no multi-page"],
        ["Verified Gross Mass",             "1,770",  "6.0%", "300 multi-page, now used in packets"],
        ["Certificate of Origin",           "1,754",  "5.9%", "Single-page only"],
        ["Shipper's Letter of Instruction", "1,721",  "5.8%", "300 multi-page, now used in packets"],
        ["Cargo Manifest",                  "1,576",  "5.3%", "500 multi-page"],
        ["Import/Export License",           "1,080",  "3.6%", "300 multi-page"],
        ["Power of Attorney",                 "436",  "1.5%", "Single-page only — underrepresented"],
        ["Customs Declaration (CN23)",        "261",  "0.9%", "Single-page only — most underrepresented"],
    ]
)
sp()
body(
    "CI and PL dominate at 47% combined because they have the most multi-page PDFs and "
    "appear in nearly every splitting template. CN23 and POA are underrepresented. "
    "Mitigation: use class-weighted loss during training (see Section 6.3)."
)

sp()
h2("3.4 Splitting Packet Design")
body(
    "5,000 synthetic combined-PDF packets simulate real DHL document batches. "
    "Each packet contains 2-13 pages (avg 5.6) with 2-7 complete documents in sequence."
)
bullet("35% of documents in packets are multi-page (was 0% before dataset rebuild)")
bullet("Templates include same-class-back-to-back scenarios (e.g. 2× CI, 2× HBL)")
bullet("Document order within each packet is randomly shuffled")
bullet("8 of 12 classes can appear as multi-page documents in packets")

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4: TRAINING PROMPT FORMAT
# ══════════════════════════════════════════════════════════════════════════
h1("4. Training Prompt Format")

h2("4.1 The Prompt — with Previous Page Context")
body(
    "Every training example includes the previous page's classification as context. "
    "This matches real inference (where you always process pages in sequence) and "
    "makes same-class document boundary detection possible."
)

sp()
code(
    "# ── USER TURN ──────────────────────────────────────────────────────\n"
    "[IMAGE: PNG of the current document page at 150 DPI]\n"
    "\n"
    "Analyze this DHL logistics document page.\n"
    "\n"
    "Previous page: <prev_label>\n"
    "\n"
    "Output format: <document_class> | <START or CONTINUATION>\n"
    "\n"
    "Document classes: Commercial Invoice, House Bill of Lading,\n"
    "Certificate of Origin, Shipper's Letter of Instruction,\n"
    "Dangerous Goods Declaration, Verified Gross Mass, House Airway Bill,\n"
    "Packing List, Customs Declaration, Cargo Manifest,\n"
    "Import/Export License, Power of Attorney\n"
    "\n"
    "START = first page of a new document\n"
    "CONTINUATION = this page continues the same document as the previous page\n"
    "\n"
    "# ── ASSISTANT TURN (ground truth label) ────────────────────────────\n"
    "Commercial Invoice | CONTINUATION"
)

sp()
h2("4.2 Values for <prev_label>")
add_table(
    ["Situation", "prev_label value", "Example"],
    [
        ["First page of the batch / standalone doc",
         "none (first page of batch)",
         "Page 1 of any packet or single doc"],
        ["Any subsequent page",
         "The exact output of the previous page's prediction",
         "'Commercial Invoice | START'\n'House Bill of Lading | CONTINUATION'"],
    ]
)

sp()
h2("4.3 What the Model Learns from Sequence Context")
add_table(
    ["prev_label", "Current page visual", "Expected output", "Why this is learnable"],
    [
        ["none (first page of batch)",    "CI header + full layout",             "Commercial Invoice | START",        "First page — always START"],
        ["Commercial Invoice | START",    "CI table rows, no header",            "Commercial Invoice | CONTINUATION", "Same doc continues — table header missing"],
        ["Commercial Invoice | CONT.",    "CI table rows, no header",            "Commercial Invoice | CONTINUATION", "Still same doc"],
        ["Commercial Invoice | CONT.",    "NEW CI header with different shipper", "Commercial Invoice | START",        "Fresh header = new document — critical boundary case!"],
        ["House Bill of Lading | START",  "CI header + full layout",             "Commercial Invoice | START",        "Class changed — always START"],
    ]
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5: INFERENCE — INPUT AND OUTPUT
# ══════════════════════════════════════════════════════════════════════════
h1("5. Inference — Input and Output")

h2("5.1 Single Page Inference")
body(
    "For a single document page (standalone or as part of a batch), the model takes "
    "one image and one text prompt, and returns one string."
)

sp()
code(
    "from transformers import Qwen2_5_VLForConditionalGeneration, AutoProcessor\n"
    "from PIL import Image\n"
    "\n"
    "model = Qwen2_5_VLForConditionalGeneration.from_pretrained('path/to/finetuned-model')\n"
    "processor = AutoProcessor.from_pretrained('path/to/finetuned-model')\n"
    "\n"
    "def classify_page(image: Image.Image, prev_label: str) -> str:\n"
    "    \"\"\"\n"
    "    Returns e.g. 'Commercial Invoice | START'\n"
    "    \"\"\"\n"
    "    prompt = (\n"
    "        'Analyze this DHL logistics document page.\\n\\n'\n"
    "        f'Previous page: {prev_label}\\n\\n'\n"
    "        'Output format: <document_class> | <START or CONTINUATION>\\n\\n'\n"
    "        'Document classes: Commercial Invoice, House Bill of Lading, '\n"
    "        'Certificate of Origin, Shipper\\'s Letter of Instruction, '\n"
    "        'Dangerous Goods Declaration, Verified Gross Mass, House Airway Bill, '\n"
    "        'Packing List, Customs Declaration, Cargo Manifest, '\n"
    "        'Import/Export License, Power of Attorney\\n\\n'\n"
    "        'START = first page of a new document\\n'\n"
    "        'CONTINUATION = this page continues the same document as the previous page'\n"
    "    )\n"
    "    messages = [{'role': 'user', 'content': [\n"
    "        {'type': 'image', 'image': image},\n"
    "        {'type': 'text',  'text': prompt}\n"
    "    ]}]\n"
    "    text = processor.apply_chat_template(messages, add_generation_prompt=True)\n"
    "    inputs = processor(text=[text], images=[image], return_tensors='pt').to(model.device)\n"
    "    out = model.generate(**inputs, max_new_tokens=20)\n"
    "    return processor.decode(out[0], skip_special_tokens=True).strip()\n"
)

sp()
body(
    "Current implementation: vision-only (image input). "
    "Future enhancement: add a text input path so the same model can classify and split "
    "from OCR-extracted text when an image is unavailable or compute must be minimised. "
    "Both paths will share the same 'Previous page:' context format and output schema."
)

sp()
h2("5.2 Processing a Full Multi-Page PDF")
body(
    "For a combined batch PDF, process pages in order and carry the previous label forward. "
    "This mirrors exactly how the training data was built."
)

sp()
code(
    "import fitz  # PyMuPDF\n"
    "from PIL import Image\n"
    "import io\n"
    "\n"
    "FIRST_PAGE_PREV = 'none (first page of batch)'\n"
    "\n"
    "def process_pdf(pdf_path: str) -> list[dict]:\n"
    "    \"\"\"\n"
    "    Returns a list of page predictions:\n"
    "    [{'page': 1, 'prediction': 'Commercial Invoice | START', 'prev': '...'},\n"
    "     {'page': 2, 'prediction': 'Commercial Invoice | CONTINUATION', 'prev': '...'},\n"
    "     ...]\n"
    "    \"\"\"\n"
    "    pdf = fitz.open(pdf_path)\n"
    "    results = []\n"
    "    prev_label = FIRST_PAGE_PREV\n"
    "\n"
    "    for page_num, page in enumerate(pdf, 1):\n"
    "        # Render page to image at 150 DPI\n"
    "        mat = fitz.Matrix(150/72, 150/72)\n"
    "        pix = page.get_pixmap(matrix=mat)\n"
    "        img = Image.open(io.BytesIO(pix.tobytes('png')))\n"
    "\n"
    "        prediction = classify_page(img, prev_label)\n"
    "        results.append({'page': page_num, 'prediction': prediction, 'prev': prev_label})\n"
    "        prev_label = prediction  # carry forward to next page\n"
    "\n"
    "    pdf.close()\n"
    "    return results\n"
    "\n"
    "# Example output for a 6-page combined PDF:\n"
    "# page 1: prev='none (first page of batch)'  → 'Commercial Invoice | START'\n"
    "# page 2: prev='Commercial Invoice | START'  → 'Commercial Invoice | CONTINUATION'\n"
    "# page 3: prev='Commercial Invoice | CONT.'  → 'House Bill of Lading | START'\n"
    "# page 4: prev='House Bill of Lading | START'→ 'House Bill of Lading | CONTINUATION'\n"
    "# page 5: prev='House Bill of Lading | CONT.'→ 'House Bill of Lading | START'  ← new HBL!\n"
    "# page 6: prev='House Bill of Lading | START'→ 'Certificate of Origin | START'\n"
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6: POST-PROCESSING — DOCUMENT SPLITTING
# ══════════════════════════════════════════════════════════════════════════
h1("6. Post-Processing — Document Splitting")

h2("6.1 Core Splitting Algorithm")
body(
    "Once the model has classified every page, splitting is a simple grouping operation. "
    "A new document begins at every START page. CONTINUATION pages belong to the most "
    "recent START. One edge case: if a CONTINUATION page's class differs from the current "
    "document's class, treat it as a new START (class mismatch overrides position label)."
)

sp()
code(
    "def split_into_documents(page_results: list[dict]) -> list[dict]:\n"
    "    \"\"\"\n"
    "    Input:  list of {'page': int, 'prediction': 'Class | START/CONTINUATION'}\n"
    "    Output: list of {'class': str, 'page_start': int, 'page_end': int, 'n_pages': int}\n"
    "\n"
    "    Rule: one document is always complete before the next begins.\n"
    "    Edge case: class mismatch on CONTINUATION → treat as new START.\n"
    "    \"\"\"\n"
    "    documents = []\n"
    "    current   = None\n"
    "\n"
    "    for item in page_results:\n"
    "        pred  = item['prediction'].strip()\n"
    "        parts = pred.split(' | ', 1)\n"
    "        cls   = parts[0] if parts else pred\n"
    "        pos   = parts[1] if len(parts) > 1 else 'START'\n"
    "        pg    = item['page']\n"
    "\n"
    "        is_new = (pos == 'START') or (current and cls != current['class'])\n"
    "\n"
    "        if is_new:\n"
    "            if current:\n"
    "                documents.append(current)          # close previous document\n"
    "            current = {'class': cls, 'page_start': pg, 'page_end': pg}\n"
    "        else:\n"
    "            if current:\n"
    "                current['page_end'] = pg           # extend current document\n"
    "            else:\n"
    "                # CONTINUATION with no prior START — treat as START\n"
    "                current = {'class': cls, 'page_start': pg, 'page_end': pg}\n"
    "\n"
    "    if current:\n"
    "        documents.append(current)                  # close last document\n"
    "\n"
    "    for d in documents:\n"
    "        d['n_pages'] = d['page_end'] - d['page_start'] + 1\n"
    "\n"
    "    return documents\n"
)

sp()
h2("6.2 Full End-to-End Example")
body(
    "Input: a 7-page combined PDF. Model predictions shown with the prev_label context."
)
sp()
code(
    "# Model predictions for a 7-page combined PDF:\n"
    "page_results = [\n"
    "    {'page': 1, 'prediction': 'Commercial Invoice | START'},\n"
    "    {'page': 2, 'prediction': 'Commercial Invoice | CONTINUATION'},\n"
    "    {'page': 3, 'prediction': 'House Bill of Lading | START'},\n"
    "    {'page': 4, 'prediction': 'House Bill of Lading | CONTINUATION'},\n"
    "    {'page': 5, 'prediction': 'House Bill of Lading | START'},   # NEW HBL!\n"
    "    {'page': 6, 'prediction': 'House Bill of Lading | CONTINUATION'},\n"
    "    {'page': 7, 'prediction': 'Certificate of Origin | START'},\n"
    "]\n"
    "\n"
    "documents = split_into_documents(page_results)\n"
    "\n"
    "# Output:\n"
    "# [{'class': 'Commercial Invoice',   'page_start': 1, 'page_end': 2, 'n_pages': 2},\n"
    "#  {'class': 'House Bill of Lading', 'page_start': 3, 'page_end': 4, 'n_pages': 2},\n"
    "#  {'class': 'House Bill of Lading', 'page_start': 5, 'page_end': 6, 'n_pages': 2},\n"
    "#  {'class': 'Certificate of Origin','page_start': 7, 'page_end': 7, 'n_pages': 1}]\n"
)

sp()
h2("6.3 Extracting Individual Document PDFs")
body(
    "Once you have the page ranges, use PyMuPDF to extract each document into its own PDF file."
)
sp()
code(
    "import fitz\n"
    "from pathlib import Path\n"
    "\n"
    "def extract_documents(source_pdf: str, documents: list[dict], out_dir: str):\n"
    "    \"\"\"\n"
    "    Writes one PDF file per identified document.\n"
    "    Filenames: 01_Commercial_Invoice_p1-2.pdf, 02_House_Bill_of_Lading_p3-4.pdf, ...\n"
    "    \"\"\"\n"
    "    Path(out_dir).mkdir(parents=True, exist_ok=True)\n"
    "    src = fitz.open(source_pdf)\n"
    "\n"
    "    for i, doc_info in enumerate(documents, 1):\n"
    "        cls    = doc_info['class'].replace('/', '-').replace(' ', '_')\n"
    "        p_from = doc_info['page_start'] - 1   # fitz uses 0-based index\n"
    "        p_to   = doc_info['page_end']   - 1\n"
    "        fname  = f\"{i:02d}_{cls}_p{doc_info['page_start']}-{doc_info['page_end']}.pdf\"\n"
    "\n"
    "        out_doc = fitz.open()\n"
    "        out_doc.insert_pdf(src, from_page=p_from, to_page=p_to)\n"
    "        out_doc.save(str(Path(out_dir) / fname))\n"
    "        out_doc.close()\n"
    "        print(f'  {fname}  ({doc_info[\"n_pages\"]} page(s))')\n"
    "\n"
    "    src.close()\n"
    "\n"
    "# Usage:\n"
    "results   = process_pdf('combined_shipment.pdf')\n"
    "documents = split_into_documents(results)\n"
    "extract_documents('combined_shipment.pdf', documents, 'output/split_docs/')\n"
    "\n"
    "# Output files:\n"
    "# output/split_docs/01_Commercial_Invoice_p1-2.pdf\n"
    "# output/split_docs/02_House_Bill_of_Lading_p3-4.pdf\n"
    "# output/split_docs/03_House_Bill_of_Lading_p5-6.pdf\n"
    "# output/split_docs/04_Certificate_of_Origin_p7-7.pdf\n"
)

sp()
h2("6.4 Complete Pipeline with Confidence Handling")
body(
    "Production use: validate predictions before acting on them. "
    "Flag low-confidence pages for human review."
)
sp()
code(
    "VALID_CLASSES = {\n"
    "    'Commercial Invoice', 'House Bill of Lading', 'Certificate of Origin',\n"
    "    \"Shipper's Letter of Instruction\", 'Dangerous Goods Declaration',\n"
    "    'Verified Gross Mass', 'House Airway Bill', 'Packing List',\n"
    "    'Customs Declaration', 'Cargo Manifest', 'Import/Export License',\n"
    "    'Power of Attorney'\n"
    "}\n"
    "\n"
    "def validate_prediction(pred: str) -> tuple[bool, str, str]:\n"
    "    \"\"\"\n"
    "    Returns (is_valid, doc_class, position).\n"
    "    Flags anything that doesn't match the expected output format.\n"
    "    \"\"\"\n"
    "    if ' | ' not in pred:\n"
    "        return False, '', ''\n"
    "    cls, pos = pred.split(' | ', 1)\n"
    "    cls = cls.strip(); pos = pos.strip()\n"
    "    if cls not in VALID_CLASSES:\n"
    "        return False, cls, pos\n"
    "    if pos not in ('START', 'CONTINUATION'):\n"
    "        return False, cls, pos\n"
    "    return True, cls, pos\n"
    "\n"
    "# In your pipeline loop:\n"
    "for item in results:\n"
    "    valid, cls, pos = validate_prediction(item['prediction'])\n"
    "    if not valid:\n"
    "        print(f\"  ⚠ Page {item['page']}: unexpected output '{item['prediction']}' — flag for review\")\n"
    "        # Fallback: treat as START of unknown class, or route to human review\n"
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7: MODEL ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
h1("7. Model Architecture & Training Config")

h2("7.1 Base Model")
add_table(
    ["Parameter", "Value", "Reason"],
    [
        ["Base model",          "Qwen3-VL-4B-Instruct",        "DocVQA 94.9 — best document understanding in the 3-5B range"],
        ["Hardware target",     "RTX 5080 16 GB GDDR7",        "QLoRA (4-bit) fits safely; LoRA (16-bit) needs 20+ GB"],
        ["Training method",     "QLoRA (4-bit) on 16 GB",      "4-bit quantization + LoRA adapters — tiny quality trade-off vs full LoRA"],
        ["Training library",    "Unsloth + TRL SFTTrainer",    "2× faster training, 60% less VRAM vs vanilla HuggingFace"],
        ["Image resolution",    "150 DPI (≈1241×1754 px A4)",  "Sufficient for all text and table structure; balanced quality/speed"],
        ["Max sequence length", "2048 tokens",                  "Image ≈ 512-768 tokens + prompt + answer ≈ ~950 total"],
        ["Est. training time",  "~45-60 min / 3 epochs",       "30K examples × 3 epochs on RTX 5080 with QLoRA"],
    ]
)

sp()
body(
    "Model benchmark comparison (DocVQA = document question answering, most relevant to this task):"
)
add_table(
    ["Model", "DocVQA", "TextVQA", "Params", "VRAM (QLoRA)", "Verdict"],
    [
        ["Qwen3-VL-4B-Instruct",   "94.9", "81.8", "4B",  "~10 GB", "✅ Chosen — best DocVQA in class"],
        ["Qwen2.5-VL-3B-Instruct", "93.9", "84.9", "3B",  "~8 GB",  "Fallback — better TextVQA, less VRAM"],
        ["Qwen2.5-VL-7B-Instruct", "95.7", "—",    "7B",  "~14 GB", "Best quality — needs 20+ GB for LoRA"],
    ]
)
body("Change the model any time by editing model.name in train_config.yaml — no code changes needed.")

sp()
h2("7.2 LoRA Configuration (all values live in train_config.yaml)")
add_table(
    ["Hyperparameter", "Value", "Reasoning"],
    [
        ["lora_r (rank)",          "32",   "Dual output (class + START/CONTINUATION) needs more capacity than r=16"],
        ["lora_alpha",             "32",   "Scale = alpha/r = 1.0 — conservative, stable for dual output tasks"],
        ["lora_dropout",           "0.05", "Light regularization; 30K examples is large enough"],
        ["target_modules",         "q,k,v,o,gate,up,down", "All attention + FFN — standard for Qwen models"],
        ["per_device_batch_size",  "1",    "RTX 5080 16 GB: batch=2 risks OOM with 4B + A4 images. Use 2 on A100/24 GB+"],
        ["gradient_accumulation",  "16",   "Effective batch = 1 × 16 = 16 (same as before, was 2 × 8)"],
        ["Learning rate",          "2e-4", "Standard for LoRA on VLMs; stable with cosine schedule"],
        ["Epochs",                 "3",    "Sweet spot for 30K examples; beyond 3 risks overfitting"],
        ["LR schedule",            "cosine + 3% warmup", "Smooth decay; warmup prevents unstable early updates"],
    ]
)

sp()
h2("7.3 train_config.yaml — Flexible Pipeline")
body(
    "All hyperparameters live in train_config.yaml. Change model, batch size, "
    "learning rate, or any setting without touching train.py."
)
code(
    "# Switch model — change one line in train_config.yaml:\n"
    "model:\n"
    "  name: 'Qwen/Qwen2.5-VL-3B-Instruct'   # or Qwen3-VL-4B, Qwen2.5-VL-7B\n"
    "  method: 'qlora'                          # qlora / lora / auto\n"
    "\n"
    "# Run with a different config file:\n"
    "python train.py --config experiments/qlora_r16.yaml\n"
    "\n"
    "# CLI flags override YAML (useful for quick experiments):\n"
    "python train.py --method lora --epochs 1 --debug\n"
)

sp()
h2("7.3 Class Weighting for Imbalanced Classes")
body(
    "CI and PL dominate training at 47% combined. CN23 has only 261 training examples. "
    "Use class weights in the loss to prevent the model from ignoring rare classes."
)
sp()
code(
    "# Approximate class weights — inverse of normalized frequency\n"
    "# Apply in the trainer via a custom compute_loss or sample weights\n"
    "CLASS_WEIGHTS = {\n"
    "    'Commercial Invoice':              1.0,   # 25.7% — baseline\n"
    "    'Packing List':                    1.2,   # 21.5%\n"
    "    'House Bill of Lading':            2.4,   # 10.7%\n"
    "    'Dangerous Goods Declaration':     3.9,   # 6.6%\n"
    "    'House Airway Bill':               3.9,   # 6.6%\n"
    "    'Verified Gross Mass':             4.3,   # 6.0%\n"
    "    'Certificate of Origin':           4.4,   # 5.9%\n"
    "    'Shippers Letter of Instruction':  4.5,   # 5.8%\n"
    "    'Cargo Manifest':                  4.9,   # 5.3%\n"
    "    'Import/Export License':           7.1,   # 3.6%\n"
    "    'Power of Attorney':              17.6,   # 1.5%\n"
    "    'Customs Declaration':            29.4,   # 0.9% — most underrepresented\n"
    "}\n"
    "# Implementation: oversample rare classes in train.jsonl by duplicating examples,\n"
    "# OR use WeightedRandomSampler in the DataLoader.\n"
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8: PROJECT STATUS & ROADMAP
# ══════════════════════════════════════════════════════════════════════════
h1("8. Project Status & Roadmap")

add_table(
    ["Phase", "Task", "Status"],
    [
        ["1", "Collect reference template PDFs (DHL/IATA/CBP)",                            "✅ Complete"],
        ["2", "Generate 12,000 single-page synthetic PDFs (12 classes, 4 variants each)",  "✅ Complete"],
        ["3", "Generate 3,300 multi-page PDFs (8 classes, 2 pages each)",                  "✅ Complete"],
        ["4", "Generate 5,000 splitting packets (combined multi-doc PDFs)",                "✅ Complete — rebuilt with 35% multi-page docs"],
        ["5", "Convert PDFs → PNGs, build train/val/test JSONL with prev-page context",    "✅ Complete — 37,093 examples, 72%/28% START/CONT"],
        ["6", "Train Qwen3-VL-4B with QLoRA via Unsloth (RTX 5080, train_config.yaml)",    "⏳ Ready — python train.py"],
        ["7", "Evaluate: per-class accuracy, START/CONT accuracy, boundary F1",            "⏳ Pending training"],
        ["8", "Fix class imbalance: generate multi-page COO/HAWB/CN23/POA",               "⏳ Recommended before or after Phase 6"],
        ["9", "Add text input path for multimodal classification + splitting",             "⏳ Future — vision-only validated first"],
        ["10", "Production pipeline: PDF → classify_page() → split_into_documents()",     "⏳ Pending Phase 6 completion"],
        ["11", "Field extraction: second fine-tuning pass for structured JSON output",     "⏳ Future work"],
    ]
)

sp()
h2("8.1 How to Run Training")
code(
    "cd D:\\finetuning\\DHL_Document_finetuning\n"
    "\n"
    "# Auto-detects VRAM — uses LoRA if >= 16 GB, QLoRA if < 16 GB\n"
    "python train.py\n"
    "\n"
    "# Force 4-bit QLoRA (lower VRAM, ~10 GB required)\n"
    "python train.py --method qlora\n"
    "\n"
    "# Smoke test — 100 examples, 1 epoch\n"
    "python train.py --epochs 1 --debug\n"
    "\n"
    "# Outputs saved to: model_output/\n"
    "# Checkpoints every 500 steps, keep last 3\n"
    "# Val eval every 500 steps\n"
)

sp()
h2("8.2 How to Regenerate the Dataset")
code(
    "cd D:\\finetuning\\DHL_Document_finetuning\n"
    "\n"
    "# Step 1 — Regenerate splitting packets (uses new templates)\n"
    "python synthetic/generate_splitting_data_v2.py --count 5000\n"
    "\n"
    "# Step 2 — Delete stale packet images (required if packets were regenerated)\n"
    "rmdir /s /q Training_Data\\images\\packets\n"
    "\n"
    "# Step 3 — Rebuild train/val/test JSONL files\n"
    "python prepare_dataset.py\n"
    "\n"
    "# To also regenerate multi-page extended docs (HBL/SLI/DGD/VGM/EEI):\n"
    "python synthetic/generate_multipage_extended.py --count 300\n"
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 9: TECHNOLOGY CHOICES
# ══════════════════════════════════════════════════════════════════════════
h1("9. Technology Choices")

add_table(
    ["Component",              "Choice",                         "Reason"],
    [
        ["Base VLM",           "Qwen2.5-VL-3B-Instruct",        "Best document understanding per VRAM; 3B fits on RTX 3090/4090"],
        ["Training framework", "Unsloth + TRL SFTTrainer",       "2× faster training, 60% less VRAM vs vanilla HuggingFace"],
        ["Fine-tuning method", "LoRA r=32 (or QLoRA 4-bit)",     "Trains only 0.2% of parameters; preserves base capabilities"],
        ["PDF → Image",        "PyMuPDF (fitz), 150 DPI",        "No poppler dependency; fast; runs natively on Windows"],
        ["PDF generation",     "ReportLab Platypus",             "Programmatic layout control; runs without GUI/browser"],
        ["Fake data",          "Faker library",                   "Realistic names, addresses, phone numbers, dates"],
        ["OCR (inference)",    "Mistral OCR 3 (future)",         "If text extraction needed alongside classification"],
        ["IDP (if needed)",    "Azure AI Document Intelligence", "If non-engineer users need to operate the pipeline"],
    ]
)

sp()
h2("9.1 Why Qwen2.5-VL-3B over Larger Models")
add_table(
    ["Model",            "Params", "VRAM (LoRA)", "Doc Understanding", "Verdict"],
    [
        ["Qwen2.5-VL-3B-Instruct",  "3B",  "~10-12 GB", "Strong",     "✅ Chosen — fits on single RTX 3090/4090"],
        ["Qwen2.5-VL-7B-Instruct",  "7B",  "~18-22 GB", "Very strong", "⚠ Needs A100 or dual GPU for LoRA"],
        ["Pixtral 12B",             "12B", "~28-32 GB", "Excellent",   "❌ Too large for single consumer GPU"],
        ["InternVL2-8B",            "8B",  "~20-24 GB", "Very strong", "⚠ Similar VRAM to Qwen2-VL-7B"],
    ]
)

pagebreak()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 10: QUICK REFERENCE — KEY FILES
# ══════════════════════════════════════════════════════════════════════════
h1("10. Quick Reference — Key Files")

add_table(
    ["File", "Purpose"],
    [
        ["synthetic/generate_all_multiformat.py",        "Master runner — regenerates all 12,000 single-page PDFs"],
        ["synthetic/generate_multipage_docs.py",         "Multi-page PDFs for CI, PL, Manifest (500 each)"],
        ["synthetic/generate_multipage_extended.py",     "Multi-page PDFs for HBL, SLI, DGD, VGM, EEI (300 each)"],
        ["synthetic/generate_splitting_data_v2.py",      "Generates 5,000 combined packet PDFs with annotations"],
        ["prepare_dataset.py",                           "Converts PDFs→PNGs, builds train/val/test JSONL with prev-page context"],
        ["train.py",                                     "LoRA/QLoRA fine-tuning via Unsloth — Qwen2.5-VL-3B"],
        ["Training_Data/train.jsonl",                    "29,744 training examples (72% START / 28% CONTINUATION)"],
        ["Training_Data/val.jsonl",                      "3,723 validation examples"],
        ["Training_Data/test.jsonl",                     "3,626 test examples"],
        ["Training_Data/dataset_stats.json",             "Counts, ratios, class distribution, source breakdown"],
        ["Training_Data/images/",                        "All PNG images: single/ multi/ packets/ subdirs"],
    ]
)

sp()
h2("10.1 Dataset Stats at a Glance")
code(
    "python -c \"\n"
    "import json\n"
    "s = json.loads(open('Training_Data/dataset_stats.json').read())\n"
    "print(f'Total: {s[\\\"total\\\"]}')\n"
    "print(f'Train/Val/Test: {s[\\\"train\\\"]}/{s[\\\"val\\\"]}/{s[\\\"test\\\"]}')\n"
    "print(f'START: {s[\\\"start_count\\\"]} | CONTINUATION: {s[\\\"continuation_count\\\"]}')\n"
    "\""
)

sp()
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.add_run(
    f"DHL Document Intelligence Fine-Tuning Reference v2.0  |  "
    f"Generated {datetime.date.today().strftime('%d %B %Y')}\n"
    "Confidential — Internal Use Only"
).font.size = Pt(8)

# ── Save ──────────────────────────────────────────────────────────────────
out_path = r"D:\finetuning\DHL_Document_finetuning\DHL_Document_Finetuning_Reference.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
