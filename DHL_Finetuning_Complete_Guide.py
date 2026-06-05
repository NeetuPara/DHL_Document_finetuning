"""Generates the DHL Fine-Tuning Complete Reference Guide DOCX."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────
DHL_RED   = RGBColor(0xD4, 0x05, 0x11)
NAVY      = RGBColor(0x1A, 0x1A, 0x2E)
TEAL      = RGBColor(0x00, 0x7B, 0x83)
ORANGE    = RGBColor(0xFF, 0x6B, 0x35)
GRAY_TEXT = RGBColor(0x44, 0x44, 0x44)
WHITE_CLR = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper functions ──────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def h1(text, color=DHL_RED):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = color; return p

def h2(text, color=NAVY):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = color; return p

def h3(text, color=TEAL):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = color; return p

def body(text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.size = Pt(size)
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + ": "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5)
    return p

def callout(text, bg="FFF3CD", border="F0A500"):
    """Highlighted callout box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade_cell(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text); run.font.size = Pt(10.5)
    return t

def info_table(headers, rows, col_widths=None, header_bg="1A1A2E"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hcells[i], header_bg)
        p = hcells[i].paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE_CLR
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        if ri % 2 == 0:
            for c in cells: shade_cell(c, "F5F5F5")
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val)); r.font.size = Pt(9.5)
    return t

def pb(): doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("DHL Document Intelligence\nFine-Tuning Complete Guide")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = DHL_RED

doc.add_paragraph()
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run(
    "A Beginner-Friendly Reference\n"
    "Covering Dataset · Model · Training · Concepts"
); r2.font.size = Pt(13); r2.font.color.rgb = GRAY_TEXT

doc.add_paragraph()
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run(f"Prepared: {datetime.date.today().strftime('%d %B %Y')}").font.size = Pt(10.5)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT ARE WE BUILDING?
# ══════════════════════════════════════════════════════════════════════════
h1("1. What Are We Building? (Big Picture)")

body(
    "Imagine DHL receives hundreds of PDFs every day from customers — invoices, packing lists, "
    "certificates, airway bills — all mixed together. A human employee would need to open each "
    "PDF, figure out what type of document it is, find the important information, and enter it "
    "into the system. That takes time and is prone to errors."
)
doc.add_paragraph()
body(
    "We are training an Artificial Intelligence model to do this automatically. The model will "
    "look at a page image of any DHL document and:"
)
bullet("CLASSIFY it — tell us 'this is a Commercial Invoice'")
bullet("SPLIT a mixed packet — 'pages 1-2 are one invoice, page 3 is a packing list, page 4 is a bill of lading'")
bullet("EXTRACT data — pull out specific fields like shipper name, invoice number, total value")

doc.add_paragraph()
callout(
    "Think of it like training a new employee who learns by studying thousands of examples. "
    "The more examples they see, the better they get at recognizing document types and reading them."
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2 — THE 12 DOCUMENT CLASSES
# ══════════════════════════════════════════════════════════════════════════
h1("2. The 12 Document Classes")

body(
    "We trained the AI on 12 types of DHL logistics documents. Each type has its own purpose, "
    "fields, and layout. Here is each class explained simply:"
)
doc.add_paragraph()

classes_info = [
    ("01", "Commercial Invoice",
     "The most important document in any international shipment. Created by the seller/shipper.",
     "Who is sending? To whom? What goods? What value? HS (tariff) codes. Currency. Incoterms.",
     "Used by customs to calculate import duty. Every international shipment needs one."),
    ("02", "House Bill of Lading (HBL)",
     "The 'receipt' for ocean freight. DHL or freight forwarder issues this to the shipper.",
     "B/L number. Vessel name. Port of loading & discharge. Container numbers. Freight terms (paid/collect).",
     "The shipper uses the B/L to claim their cargo at the destination port."),
    ("03", "Certificate of Origin (COO)",
     "Proves where the goods were manufactured. Required for preferential customs duty under trade agreements.",
     "Exporter details. Goods description. HS code. Origin criterion (A/B/C/D). Country of origin.",
     "Without a COO, goods may be taxed at the full rate instead of a reduced trade-agreement rate."),
    ("04", "Shipper's Letter of Instruction (SLI)",
     "A form the shipper fills out to tell DHL how to handle the export. Legally authorizes DHL to act.",
     "Shipper (USPPI) details. Consignee. Commodity codes (Schedule B). Export license info. Incoterms.",
     "DHL needs this to file export declarations on behalf of the shipper in the US system."),
    ("05", "Dangerous Goods Declaration (DGD)",
     "Required whenever hazardous materials are shipped by air or sea. Only trained people can complete it.",
     "UN number. Proper shipping name. Hazard class. Packing group. Quantity. Emergency contact.",
     "Airlines can refuse shipment or face heavy fines if DG declaration is missing or incorrect."),
    ("06", "Verified Gross Mass (VGM)",
     "The verified weight of a packed container. Mandatory since 2016 (SOLAS regulation) before loading onto a ship.",
     "Container number. Tare weight. Cargo weight. Total VGM. Method used (1=weigh container, 2=sum of items).",
     "If container weight is wrong, it can destabilize a ship — this law prevents accidents at sea."),
    ("07", "House Airway Bill (HAWB)",
     "The 'receipt' for air freight. Issued by DHL/freight forwarder to the shipper.",
     "HAWB number. Airports (departure & destination). Weight. Chargeable weight. Rate. Total charges.",
     "Like a B/L but for air. The shipper uses it to track and claim their cargo."),
    ("08", "Packing List",
     "Detailed list of every item in every package in the shipment.",
     "Package marks. Number of packages. Description. Net weight. Gross weight. Dimensions. CBM (volume).",
     "Used by customs to check contents match the invoice. Used by receiver to verify delivery."),
    ("09", "Customs Declaration (CN23/CN22)",
     "Postal/parcel customs form attached to the outside of international parcels.",
     "Sender name & address. Addressee. Category (gift/sale/sample). Item descriptions. Value. HS code.",
     "Required by every postal service for international parcels — postal equivalent of a commercial invoice."),
    ("10", "Cargo Manifest",
     "A summary list of ALL cargo on a specific vessel or aircraft, prepared by the carrier.",
     "Vessel or flight. All AWB/B/L numbers. Each shipper & consignee. Weight per shipment. Total cargo.",
     "Customs authorities require the manifest before the vessel/aircraft arrives to plan inspections."),
    ("11", "Import/Export License (EEI/CBP 7501)",
     "Government forms for controlled goods. In the US: CBP Form 7501 for imports, EEI for exports.",
     "Entry number. Importer/exporter. HTS codes. Country of origin. Entered value. Duty rate. Duty amount.",
     "Required for controlled goods (technology, weapons-related, dual-use items) or for customs entry."),
    ("12", "Power of Attorney (POA)",
     "A legal document where a company gives DHL the authority to act as their customs agent.",
     "Grantor (the company). Grantee (DHL). Scope (export/import/both). EIN/tax ID. Notary signature.",
     "Without a POA, DHL cannot legally file customs declarations on behalf of the customer."),
]

for num, name, what, fields, why in classes_info:
    h3(f"Class {num}: {name}")
    body("What it is: ", bold=True, size=10.5)
    doc.paragraphs[-1].add_run(what).font.size = Pt(10.5)
    body("Key fields: ", bold=True, size=10.5)
    doc.paragraphs[-1].add_run(fields).font.size = Pt(10.5)
    body("Why it matters: ", bold=True, size=10.5)
    doc.paragraphs[-1].add_run(why).font.size = Pt(10.5)
    doc.add_paragraph()

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE DATASET
# ══════════════════════════════════════════════════════════════════════════
h1("3. Our Synthetic Dataset — How It Was Built")

h2("3.1 Why Synthetic Data?")
body(
    "We cannot use real DHL customer documents for training because they contain sensitive "
    "personal and business information (names, addresses, shipment values). Instead, we "
    "generated synthetic (fake but realistic) documents using Python."
)
bullet("All company names, addresses, and values are randomly generated — they look real but belong to no one.")
bullet("All 12 document classes have the correct field structure, table layouts, and terminology.")
bullet("We created multiple visual styles per class to mimic the real world where every company uses a different template.")
doc.add_paragraph()

h2("3.2 Three Datasets — Three Jobs")
body("We created three separate datasets, each serving a different purpose:")
doc.add_paragraph()

info_table(
    ["Dataset", "Purpose", "PDFs", "What's Inside"],
    [
        ["Synthetic_Data/", "Classification + Extraction training",
         "12,000", "One document per PDF (1 page). 1,000 per class. 3-4 visual format variants."],
        ["Synthetic_Data_MultiPage/", "Teaching CONTINUATION pages",
         "3,000", "Two-page documents. Page 1 = first page, Page 2 = continuation of same document."],
        ["Synthetic_Data_Splitting_v2/", "Splitting task training",
         "5,000", "Multi-document packets. Each PDF has 2-10 pages of DIFFERENT document types mixed together."],
    ]
)

doc.add_paragraph()
h2("3.3 Format Variants — Why Multiple Layouts?")
body(
    "In real life, every company creates Commercial Invoices differently. "
    "A small exporter might use a simple Word template. A large corporation uses SAP software "
    "which generates a completely different layout. DHL has its own template. A freight "
    "forwarder has yet another format."
)
doc.add_paragraph()
body(
    "If we train the model on ONLY one layout, it will fail on the others. "
    "So for each of the 12 classes, we created 3-4 visually different formats:"
)
doc.add_paragraph()
info_table(
    ["Class", "Format 1", "Format 2", "Format 3", "Format 4"],
    [
        ["Commercial Invoice", "DHL Express (red header)", "Corporate/SAP (bank details, duty)", "Freight Forwarder (marks & numbers)", "E-commerce/ERP (SKU-based)"],
        ["House B/L", "Standard Ocean HBL", "FIATA Multimodal", "Short-Form BOL", "-"],
        ["Certificate of Origin", "Chamber of Commerce", "FTA/USMCA Numbered", "EUR.1 Movement Cert", "-"],
        ["All 12 classes", "3-4 distinct visual styles", "Different color schemes", "Different table structures", "Different terminology"],
    ]
)
doc.add_paragraph()
callout(
    "Key insight: The AI must learn the CONCEPT of 'Commercial Invoice' — not just one specific layout. "
    "Multiple format variants force the model to learn what makes a CI a CI, regardless of how it looks visually."
)

pb()
h2("3.4 Class-by-Class Dataset Details")

class_detail = [
    ("01 Commercial Invoice", "1,000 single-doc", "500 2-page (multi)", "100% in splitting", "4", "1-10 line items, 10 currencies, 24 countries, all Incoterms"),
    ("02 House Bill of Lading", "1,000 single-doc", "300 2-page (multi)", "32.3% in splitting", "3", "10-40 containers, global ports, vessel names"),
    ("03 Certificate of Origin", "1,000 single-doc", "No multi", "22.7% in splitting", "3", "FTA types (USMCA/CPTPP/EUR.1), origin criteria A-D"),
    ("04 SLI", "1,000 single-doc", "300 2-page (multi)", "26.8% in splitting", "3", "18-40 Schedule B items, EEI/ECCN codes, all export modes"),
    ("05 DGD", "1,000 single-doc", "300 2-page (multi)", "25.5% in splitting", "3", "10 UN numbers, IATA/IMDG formats, non-DG declarations"),
    ("06 VGM", "1,000 single-doc", "300 2-page (multi)", "25.7% in splitting", "3", "Method 1 & 2, 35-50 containers for multi-page"),
    ("07 HAWB", "1,000 single-doc", "No multi", "41.6% in splitting", "3", "IATA standard, Express, Consolidated styles"),
    ("08 Packing List", "1,000 single-doc", "500 2-page (multi)", "58.8% in splitting", "3", "40-60 items for multi-page, warehouse/e-commerce styles"),
    ("09 Customs Declaration", "1,000 single-doc", "No multi", "31.4% in splitting", "3", "CN22, CN23, US CBP informal entry"),
    ("10 Cargo Manifest", "1,000 single-doc", "500 2-page (multi)", "24.2% in splitting", "3", "Air & Ocean, 45-65 entries for multi-page"),
    ("11 Import/Export License", "1,000 single-doc", "300 2-page (multi)", "33.5% in splitting", "3", "CBP 7501, EEI filing, formal license document"),
    ("12 Power of Attorney", "1,000 single-doc", "No multi", "31.1% in splitting", "3", "Export/Import/Both scopes, legal & simplified styles"),
]
info_table(
    ["Class", "Single-Doc", "Multi-Page Pool", "Splitting Coverage", "Formats", "Diversity"],
    class_detail
)

pb()
h2("3.5 The Splitting Dataset — How It Works")
body(
    "Real DHL document packets often arrive as one big PDF with multiple documents inside. "
    "The splitting dataset teaches the model to find document boundaries."
)
doc.add_paragraph()
body("Example of a 7-page packet (packet_0002.pdf):", bold=True)
doc.add_paragraph()
info_table(
    ["Page", "What's on the page", "Label the model should output", "Why"],
    [
        ["Page 1", "A Commercial Invoice — first page (has company header, invoice title)", "Commercial Invoice | START", "It's the beginning of a new document"],
        ["Page 2", "SAME Commercial Invoice — continuation (table rows continue, no header)", "Commercial Invoice | CONTINUATION", "Same invoice, page 2 is a continuation"],
        ["Page 3", "A DIFFERENT Commercial Invoice (different company, different invoice#)", "Commercial Invoice | START", "New document! Even though same class"],
        ["Page 4", "Packing List (new document begins)", "Packing List | START", "Class changed = boundary"],
        ["Page 5", "Same Packing List continues (more items)", "Packing List | CONTINUATION", "Same PL, just continues"],
        ["Page 6", "Certificate of Origin", "Certificate of Origin | START", "New document boundary"],
        ["Page 7", "House Airway Bill", "House Airway Bill | START", "New document boundary"],
    ]
)
doc.add_paragraph()
callout(
    "Important: Two pages with the SAME class does NOT always mean same document. "
    "The model must distinguish between 'same document continuation' and 'new document of same type'."
    " It learns this from visual cues: full headers = START, table-only continuation = CONTINUATION."
)
doc.add_paragraph()

info_table(
    ["Shipment Type", "Packets", "Typical Documents Inside"],
    [
        ["Ocean Freight (25%)", "1,250", "CI + HBL + PL + COO + VGM + DGD + SLI + EEI + POA + Manifest"],
        ["Air Freight (25%)", "1,250", "CI + HAWB + PL + COO + DGD + SLI + EEI + POA"],
        ["Express/Postal (25%)", "1,250", "CI + CN23 (x1-4 parcels) + HAWB + POA"],
        ["Mixed/Complex (25%)", "1,250", "Random combination of all 12 classes (equal weights)"],
    ]
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4 — TRAINING DATA FORMAT
# ══════════════════════════════════════════════════════════════════════════
h1("4. Training Data Format — What the AI Sees")

h2("4.1 From PDF to Image")
body(
    "The AI model we are using is a Vision Language Model (VLM) — it processes images, not PDFs. "
    "So every PDF page is converted to a PNG image at 150 DPI (dots per inch). "
    "150 DPI gives a clear image (1,241 x 1,754 pixels for A4) without being too large."
)
doc.add_paragraph()
info_table(
    ["Source", "Images Generated", "How"],
    [
        ["Synthetic_Data/ (12,000 PDFs × 1 page)", "12,000 images", "Page 1 of each single-doc PDF"],
        ["Synthetic_Data_MultiPage/ (3,000 × 2 pages)", "6,000 images", "Both pages of each 2-page doc"],
        ["Synthetic_Data_Splitting_v2/ (5,000 × avg 5.3 pages)", "~26,500 images", "Every page of every packet"],
        ["TOTAL", "~44,500 images", "Each image = one training example"],
    ]
)

doc.add_paragraph()
h2("4.2 The Training Instruction")
body(
    "Each training example consists of three things: an image, an instruction (the question "
    "we ask the model), and the expected answer. Together these form a conversation:"
)
doc.add_paragraph()

callout(
    "INSTRUCTION (what we ask the model):\n\n"
    "'Analyze this DHL logistics document page.\n\n"
    "Output format: <document_class> | <START or CONTINUATION>\n\n"
    "Document classes: Commercial Invoice, House Bill of Lading, Certificate of Origin, "
    "Shipper's Letter of Instruction, Dangerous Goods Declaration, Verified Gross Mass, "
    "House Airway Bill, Packing List, Customs Declaration, Cargo Manifest, "
    "Import/Export License, Power of Attorney\n\n"
    "START = first page of a new document\n"
    "CONTINUATION = this page continues the same document as the previous page'"
)
doc.add_paragraph()
body("EXPECTED ANSWER (what the model must learn to output):", bold=True)
body("Commercial Invoice | START")
doc.add_paragraph()
body("Or for a continuation page:")
body("Commercial Invoice | CONTINUATION")

doc.add_paragraph()
h2("4.3 Understanding Every Part of the Instruction")
body("Let's break down what each piece of the instruction means:")
doc.add_paragraph()

info_table(
    ["Part of Instruction", "What it means", "Why it's there"],
    [
        ["'Analyze this DHL logistics document page'",
         "Tells the AI what kind of task this is and what it's looking at",
         "Sets context. The AI needs to know this is a shipping document, not a newspaper or medical form."],
        ["'Output format: <document_class> | <START or CONTINUATION>'",
         "Tells the AI exactly how to structure its answer — class name, then pipe symbol, then position",
         "If we don't specify format, different models might answer in different ways (one says 'It is a CI', another says 'CI'). Structured format makes parsing easy."],
        ["The list of 12 class names",
         "These are the only valid answers for the class part of the output",
         "Constrains the AI to pick from our 12 classes only. Without this, it might invent classes like 'shipping document' or 'logistics form'."],
        ["'START = first page of a new document'",
         "Defines what START means",
         "The AI needs to know the definition, not just the word. It will look for visual signals: company header, document title, field labels at top."],
        ["'CONTINUATION = this page continues the same document as the previous page'",
         "Defines CONTINUATION",
         "The AI learns to look for continuation signals: table rows that start mid-column, 'Page 2 of 3' footer, no document title at top."],
        ["The pipe symbol | between class and position",
         "A separator character between the two parts of the answer",
         "Makes it trivial to split the answer into class and position. We just split on ' | '."],
    ]
)

doc.add_paragraph()
h2("4.4 Why One Instruction for Both Tasks?")
body(
    "You might wonder: why do we use the SAME instruction for classification AND splitting? "
    "The answer is elegant — both tasks are actually the same operation at the page level:"
)
bullet("Classification: give it one page → outputs the class. Done.")
bullet("Splitting: give it each page of a packet one by one → outputs class + START/CONTINUATION for each → wherever START appears = document boundary found.")
doc.add_paragraph()
callout(
    "The model does not need to see multiple pages at once to find boundaries. "
    "It detects boundaries by recognizing visual signals on each individual page: "
    "a full company header means START, a table continuation without a header means CONTINUATION."
)

pb()
h2("4.5 Train / Validation / Test Split")
body("We split all 44,500 training examples into three groups:")
doc.add_paragraph()

info_table(
    ["Split", "% of data", "Examples", "Purpose"],
    [
        ["TRAIN", "80%", "~35,600", "The model actually learns from these. It sees them repeatedly during training."],
        ["VALIDATION (VAL)", "10%", "~4,450", "Checked during training to see if the model is improving or overfitting. NOT used for learning."],
        ["TEST", "10%", "~4,450", "Held out completely until AFTER training. Used for the final honest evaluation of model performance."],
    ]
)
doc.add_paragraph()
body("How we split:", bold=True)
body(
    "We split by SOURCE DOCUMENT (PDF), not by individual page. This is critical. "
    "If page 1 and page 2 of the same PDF go into different splits, the model could 'cheat' — "
    "it memorizes specific documents rather than learning general patterns."
)
bullet("All pages from the same PDF always go into the SAME split (train, val, or test).")
bullet("This is called a 'document-level split' and gives an honest picture of model performance.")

doc.add_paragraph()
info_table(
    ["Label Type", "Total", "In Train", "In Val", "In Test"],
    [
        ["START", "~35,000", "~28,000", "~3,500", "~3,500"],
        ["CONTINUATION", "~9,500", "~7,600", "~950", "~950"],
        ["TOTAL", "~44,500", "~35,600", "~4,450", "~4,450"],
    ]
)
doc.add_paragraph()
callout(
    "Notice: CONTINUATION is only ~21% of total examples. This imbalance exists because "
    "most documents are single-page (only START). During training, we may need to "
    "oversample CONTINUATION examples so the model learns that case equally well."
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5 — THE MODEL
# ══════════════════════════════════════════════════════════════════════════
h1("5. Which Model Do We Use and Why?")

h2("5.1 What is a Vision Language Model (VLM)?")
body(
    "A Vision Language Model is an AI model that can understand both images and text. "
    "Unlike a regular language model (like GPT) which only reads text, a VLM can look at "
    "an image AND answer questions about it in text."
)
doc.add_paragraph()
info_table(
    ["Model Type", "Input", "Output", "Example"],
    [
        ["Language Model (LLM)", "Text only", "Text", "ChatGPT reading a document's text"],
        ["Vision Model (CV)", "Image only", "Classification/detection", "Image classifier saying 'cat' or 'dog'"],
        ["Vision Language Model (VLM)", "Image + Text instruction", "Text", "Our model: sees document image + instruction, outputs class label"],
    ]
)
doc.add_paragraph()

h2("5.2 Why 3B is Enough — One Page at a Time")
body(
    "Before choosing a model, an important question was asked: "
    "do we actually NEED a large 7B model? The answer is NO — and here is why."
)
doc.add_paragraph()
callout(
    "Key insight: At inference time, the model receives EXACTLY ONE page image and outputs "
    "a short label (~25 characters). It does not need to read an entire document, write "
    "paragraphs, or reason through complex logic. This is a narrow, well-defined task. "
    "A smaller model fine-tuned on our specific data beats a larger general model every time."
)
doc.add_paragraph()
info_table(
    ["What our model needs to do", "Does 3B handle it?", "Does 7B handle it better?"],
    [
        ["Look at one page image (1241 x 1754 px)", "Yes — any VLM can do this", "No real difference"],
        ["Recognize document title at top of page", "Yes — basic OCR", "Slightly better on degraded scans"],
        ["Read HS codes, invoice numbers in small table text", "Yes — Qwen2.5-VL is OCR-specialist", "Marginally"],
        ["Output one of 12 fixed class names", "Yes — simple classification", "No advantage"],
        ["Decide START vs CONTINUATION from visual layout", "Yes — learned during fine-tuning", "No advantage"],
        ["Handle edge cases / unusual formats", "Good after fine-tuning", "Slightly better zero-shot"],
    ]
)
doc.add_paragraph()

h2("5.3 Our Chosen Model: Qwen2.5-VL-3B-Instruct")
body("We chose Qwen2.5-VL-3B-Instruct (the 2.5 series, released late 2024). Here is why:")
doc.add_paragraph()

info_table(
    ["Property", "Detail", "Why it matters"],
    [
        ["Name", "Qwen2.5-VL-3B-Instruct", "Made by Alibaba Cloud Qwen team. Version 2.5 is a major upgrade over 2.0."],
        ["Size", "3 Billion parameters (not 7B)", "Processes one page at a time — 3B is more than sufficient. Faster + cheaper."],
        ["Type", "Vision Language Model — document specialist", "Specifically designed for documents, forms, tables, receipts, invoices."],
        ["OCR capability", "Best-in-class for its size", "Can accurately read small text: HS codes, invoice numbers, field labels in tables."],
        ["Fine-tuning support", "Yes — unsloth has native Qwen2.5-VL support", "Simple QLoRA fine-tuning setup with proven community examples."],
        ["License", "Apache 2.0 (commercial use allowed)", "DHL can deploy this in production freely."],
        ["Hardware needed", "~8 GB VRAM for inference, ~8 GB for QLoRA fine-tuning", "Runs on a single RTX 3090/4090. No expensive A100 required."],
        ["Speed", "~0.3 seconds per page", "Processes 200 pages/minute — practical for production volume."],
    ]
)

doc.add_paragraph()
h2("5.4 Document Benchmark Comparison")
body(
    "The best way to compare models for our task is to look at document-specific benchmarks. "
    "These tests measure how well a model can read and understand document images — "
    "directly relevant to DHL document classification."
)
doc.add_paragraph()
body("What these benchmarks test:", bold=True)
bullet("DocVQA — Answer questions about document images (forms, invoices, reports). Most relevant to our task.")
bullet("TextVQA — Read text accurately from scene images and answer questions.")
bullet("ChartQA — Understand charts and structured visual data.")
doc.add_paragraph()

info_table(
    ["Model", "Size", "DocVQA", "TextVQA", "ChartQA", "Our Task Fit", "Decision"],
    [
        ["Qwen2.5-VL-3B-Instruct", "3B", "~84%", "~79%", "~74%", "Excellent — document specialist", "CHOSEN"],
        ["Qwen2.5-VL-7B-Instruct", "7B", "~91%", "~84%", "~83%", "Excellent — but overkill for classification", "Backup if 3B insufficient"],
        ["Gemma 3 4B-IT (Google)", "4B", "~69%", "~68%", "~62%", "General VLM — not doc-specialized", "Rejected for doc tasks"],
        ["InternVL2-4B", "4B", "~78%", "~73%", "~71%", "Good alternative", "Alternative if Qwen unavailable"],
        ["GPT-4V (OpenAI)", "Unknown", "~88%", "~78%", "~78%", "Strong but closed, no fine-tuning", "Rejected — privacy + cost"],
    ]
)
doc.add_paragraph()
callout(
    "Qwen2.5-VL-3B scores 15-20 percentage points HIGHER than Gemma 3 4B on document tasks "
    "(DocVQA: 84% vs 69%). This gap exists BEFORE fine-tuning. After fine-tuning on our "
    "44,500 DHL-specific labeled images, Qwen2.5-VL-3B will push well above 90% accuracy "
    "on our narrow task. Gemma 3 is an excellent general model but was not built for documents."
)

doc.add_paragraph()
h2("5.5 Qwen2.5-VL vs Qwen2-VL — What Changed?")
body(
    "The 2.5 version (late 2024) is a significant improvement over 2.0. Key upgrades:"
)
info_table(
    ["Capability", "Qwen2-VL-7B (old)", "Qwen2.5-VL-3B (new, chosen)"],
    [
        ["Document OCR accuracy", "Good", "Significantly better — especially on dense tables and small fonts"],
        ["Native image resolution", "Fixed resolution", "Dynamic resolution — adapts to image size, preserves detail"],
        ["Instruction following", "Good", "More reliable structured output (our | format works better)"],
        ["Fine-tuning stability", "Good", "Better — fewer training instabilities with QLoRA"],
        ["Parameters needed", "7B for good quality", "3B now matches old 7B quality on document tasks"],
    ]
)

doc.add_paragraph()
h2("5.6 Why Not Use a Bigger API Model (GPT-4V, Gemini)?")
bullet("Cannot be fine-tuned — we cannot train them on our DHL-specific data.", bold_prefix="No fine-tuning")
bullet("Every page sent for classification costs money per API call at production scale.", bold_prefix="Expensive at scale")
bullet("DHL documents contain sensitive shipper/consignee data — must not leave your infrastructure.", bold_prefix="Data privacy")
bullet("Dependent on external service uptime — DHL operations cannot stop if API goes down.", bold_prefix="Reliability")
doc.add_paragraph()
body(
    "By fine-tuning Qwen2.5-VL-3B locally, we get a model that is DHL-specific, private, "
    "free to run after setup, and faster than any external API for batch processing."
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6 — WHAT IS FINE-TUNING?
# ══════════════════════════════════════════════════════════════════════════
h1("6. What is Fine-Tuning? (Concept Explained Simply)")

h2("6.1 The Analogy")
callout(
    "Imagine hiring a brilliant university graduate who has read millions of books and can answer "
    "questions on almost any topic. But they have never worked in logistics before. Fine-tuning "
    "is like giving them an intensive 2-week training course on DHL documents. After the course, "
    "they are still the same smart person — but now they are also an expert at reading DHL invoices, "
    "packing lists, and bills of lading."
)
doc.add_paragraph()

h2("6.2 Pre-training vs Fine-tuning")
info_table(
    ["Stage", "What happens", "Data used", "Time", "Who does it"],
    [
        ["Pre-training",
         "Model learns language, vision, general knowledge from the internet. Learns what words, images, and concepts mean.",
         "Billions of web pages, books, images",
         "Months on thousands of GPUs",
         "Qwen team at Alibaba"],
        ["Fine-tuning (our job)",
         "We show the model 44,500 DHL document page images with correct labels. It adjusts its LoRA adapters to become a DHL document specialist.",
         "Our 44,500 labeled page images (12 classes, START/CONTINUATION labels)",
         "3-5 hours on 1 GPU (3B model is fast to fine-tune)",
         "Us / DHL team"],
    ]
)

doc.add_paragraph()
h2("6.3 What is LoRA? (Low-Rank Adaptation)")
body(
    "Even a 3 billion parameter model has 3,000,000,000 numbers inside it. "
    "Fine-tuning ALL 3 billion parameters still requires significant memory. "
    "LoRA is a smart trick that makes fine-tuning much cheaper and faster."
)
doc.add_paragraph()
body("The key idea:", bold=True)
bullet("Instead of changing all 3 billion original parameters, we ADD a small set of NEW parameters (about 0.1% of the original — roughly 3 million extra numbers).")
bullet("We only train these small added parameters, keeping the original model completely frozen.")
bullet("The result: 90%+ less memory needed, 5-10x faster training, essentially same quality improvement.")
doc.add_paragraph()
info_table(
    ["Fine-tuning Method", "GPU Memory", "Training Time", "Quality"],
    [
        ["Full Fine-tuning (all 3B parameters)", "~30 GB+ VRAM", "6-10 hours", "Best possible — but impractical"],
        ["LoRA (3B + LoRA adapters, 16-bit)", "~12-16 GB (1 RTX 3090)", "3-5 hours", "Excellent — very close to full fine-tuning"],
        ["QLoRA (3B quantized to 4-bit + LoRA)", "~8-10 GB (1 RTX 3080)", "4-6 hours", "Great — tiny quality trade-off, very accessible hardware"],
    ]
)
doc.add_paragraph()
callout(
    "We will use QLoRA — 4-bit quantization + LoRA. This means we compress the model's "
    "memory usage by representing numbers with less precision (4 bits instead of 16 bits). "
    "We then add LoRA adapters on top. Result: fine-tuning a 7B model on a single GPU with 16GB RAM."
)

pb()
h2("6.4 What Happens During Fine-tuning? (Step by Step)")
body("Here is what happens during one training step, in plain English:")
doc.add_paragraph()

steps = [
    ("Step 1: Feed", "The model is shown ONE training example: an image of a document page."),
    ("Step 2: Instruction", "The model reads the instruction: 'Analyze this DHL document page. Output: <class> | <START or CONTINUATION>'"),
    ("Step 3: Predict", "The model generates its best guess. Example: it might guess 'Packing List | START'"),
    ("Step 4: Compare", "We compare the model's guess to the correct answer ('Commercial Invoice | START'). This difference is called the 'loss'."),
    ("Step 5: Learn", "The model's parameters are adjusted slightly to make it less likely to make this mistake again. The learning rate controls how big the adjustment is."),
    ("Step 6: Repeat", "This process repeats for all 35,600 training examples. One full pass through all examples is called an 'epoch'. We typically train for 2-3 epochs."),
]

for step, desc in steps:
    bullet(desc, bold_prefix=step)

doc.add_paragraph()
info_table(
    ["Training Concept", "Simple Explanation", "Technical term"],
    [
        ["How much the model adjusts each step", "Small = safe, large = risky/fast", "Learning rate (e.g. 2e-4)"],
        ["How many examples shown at once", "Larger batches = smoother learning but more memory", "Batch size (e.g. 4-8)"],
        ["How many times we go through all data", "More epochs = more learning, but risk of memorizing", "Epochs (e.g. 2-3)"],
        ["How wrong the model's answer is", "We minimize this number during training", "Loss"],
        ["Checking model on unseen data during training", "Catches if model is memorizing vs truly learning", "Validation loss"],
    ]
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7 — INSTRUCTION TUNING EXPLAINED
# ══════════════════════════════════════════════════════════════════════════
h1("7. Instruction Tuning — Teaching the Model to Follow Instructions")

h2("7.1 What is Instruction Tuning?")
body(
    "Instruction tuning is a specific type of fine-tuning where we teach the model to follow "
    "INSTRUCTIONS — not just predict the next word. We format every training example as a "
    "conversation between a user and an AI assistant."
)
doc.add_paragraph()
callout(
    "Without instruction tuning: Model sees text and tries to complete it naturally.\n"
    "With instruction tuning: Model learns to READ an instruction, LOOK at an image, and RESPOND helpfully.\n\n"
    "This is the difference between a language model and an AI assistant."
)

doc.add_paragraph()
h2("7.2 The Conversation Format")
body(
    "Every training example is structured like a chat conversation in JSON format. "
    "Here is what one example looks like and what each piece means:"
)
doc.add_paragraph()

info_table(
    ["JSON Field", "Value", "What it means"],
    [
        ['"role": "user"', 'The person asking the question', 'In our case, this is DHL\'s system submitting a document page for analysis'],
        ['"type": "image"', 'An image input', 'The PNG image of one document page goes here'],
        ['"type": "text"', 'The instruction text', 'The analysis instruction we want the model to follow'],
        ['"role": "assistant"', 'The AI\'s response', 'What the CORRECT answer should be'],
        ['"content": "Commercial Invoice | START"', 'The expected output', 'The model learns to produce exactly this format'],
        ['"image": "images/single/..."', 'Path to the PNG file', 'Where the image file is stored on disk'],
        ['"label": "Commercial Invoice | START"', 'The ground truth label', 'Used for evaluation metrics after training'],
        ['"source": "single_doc"', 'Which dataset it came from', 'Useful for analyzing performance by source type'],
    ]
)

doc.add_paragraph()
h2("7.3 Why the Pipe Symbol ( | ) in the Output?")
body(
    "The format 'Commercial Invoice | START' uses a pipe symbol to separate two pieces of information. "
    "This is not random — it is a deliberate design choice:"
)
bullet("Easy to parse: split on ' | ' gives ['Commercial Invoice', 'START']")
bullet("Unambiguous: none of the class names contain ' | ' so there is no confusion")
bullet("Compact: short output means faster generation and clearer learning signal")
bullet("Structured: models trained on structured outputs are more reliable than free-form answers")

doc.add_paragraph()
h2("7.4 What Visual Signals Does the Model Learn?")
body("The model learns to look for these visual signals when deciding START vs CONTINUATION:")
doc.add_paragraph()

info_table(
    ["Visual Signal", "Indicates", "Example"],
    [
        ["Full company letterhead / logo area at top", "START", "First page of any document has company name, address, document title at the top"],
        ["Document title (e.g. 'COMMERCIAL INVOICE')", "START", "Title only appears on first page"],
        ["Table header row (column labels like 'Description | Qty | Value')", "START (new doc) or CONTINUATION (repeated header)", "Repeated headers on page 2 of a long table = continuation"],
        ["Table rows starting mid-way (no header above)", "CONTINUATION", "Page 2 of a packing list starts right with item rows"],
        ["'Page 2 of 3' or similar footer", "CONTINUATION", "Explicit page number indicating not first page"],
        ["New company name / different invoice number visible", "START", "Different company = definitely a new document"],
        ["Signature / stamp area (usually at bottom of last page)", "Approaching END of document", "Helps detect document completion"],
        ["Blank rows at bottom of table (filler rows)", "START of a document form", "Empty form rows appear on single-page docs"],
    ]
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8 — EVALUATION
# ══════════════════════════════════════════════════════════════════════════
h1("8. How Do We Know If the Model is Good?")

h2("8.1 Classification Metrics")
body("After training, we run the model on the TEST set (never seen during training) and measure:")
doc.add_paragraph()

info_table(
    ["Metric", "What it means", "Target"],
    [
        ["Accuracy", "What % of pages did the model classify correctly?", "> 95%"],
        ["Precision (per class)", "When model says 'Commercial Invoice', how often is it right?", "> 90% per class"],
        ["Recall (per class)", "Of all actual Commercial Invoices, what % did the model find?", "> 90% per class"],
        ["F1 Score", "Balanced measure combining precision and recall", "> 90% per class"],
        ["Confusion Matrix", "Table showing which classes get confused with which others", "Few off-diagonal entries"],
    ]
)

doc.add_paragraph()
h2("8.2 Splitting Metrics")
info_table(
    ["Metric", "What it means", "Target"],
    [
        ["Boundary Accuracy", "What % of document boundaries did it find correctly?", "> 90%"],
        ["Document-level F1", "Precision & recall of identified document blocks", "> 85%"],
        ["START Accuracy", "Of all START pages, what % did it label correctly?", "> 90%"],
        ["CONTINUATION Accuracy", "Of all CONTINUATION pages, what % did it label correctly?", "> 85%"],
    ]
)

doc.add_paragraph()
h2("8.3 Common Failure Modes to Watch For")
bullet("Confusing HBL with HAWB: both are transport documents. HBL = ocean, HAWB = air.")
bullet("Confusing CI with PL: invoice has values/prices, packing list has dimensions/weights.")
bullet("Missing CONTINUATION: if model always says START, it will split every page as a new document.")
bullet("Rare classes getting lower accuracy: CN23, VGM, POA have less training data than CI or PL.")

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 9 — FULL ROADMAP
# ══════════════════════════════════════════════════════════════════════════
h1("9. The Full Fine-Tuning Roadmap")

info_table(
    ["Phase", "What we do", "Status", "Output"],
    [
        ["Phase 1", "Download reference DHL document templates for all 12 classes", "DONE", "Documents/ folder"],
        ["Phase 2A", "Generate 12,000 synthetic single-doc PDFs (12 classes x 1,000, 3-4 formats each)", "DONE", "Synthetic_Data/"],
        ["Phase 2B", "Generate 3,000 two-page continuation docs (8 classes)", "DONE", "Synthetic_Data_MultiPage/"],
        ["Phase 2C", "Generate 5,000 splitting packets (mixed docs, 2-10 pages, balanced across all 12 classes)", "DONE", "Synthetic_Data_Splitting_v2/"],
        ["Phase 2D", "Verify all datasets: counts, page integrity, annotations, label logic", "DONE", "Verification report"],
        ["Phase 3", "Convert all PDFs to PNG images (44,500 pages)", "IN PROGRESS", "Training_Data/images/"],
        ["Phase 4", "Generate train/val/test JSONL files with instruction + label per image", "NEXT", "train.jsonl, val.jsonl, test.jsonl"],
        ["Phase 5", "Set up Qwen2-VL-7B with QLoRA fine-tuning (unsloth + TRL)", "PENDING", "Training environment"],
        ["Phase 6", "Fine-tune for 2-3 epochs, monitor validation loss", "PENDING", "Fine-tuned model checkpoint"],
        ["Phase 7", "Evaluate on test set: accuracy, F1 per class, confusion matrix", "PENDING", "Evaluation report"],
        ["Phase 8", "Build inference pipeline: PDF -> pages -> classify each -> merge boundaries", "PENDING", "Production pipeline"],
    ]
)

pb()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 10 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════
h1("10. Quick Reference — Numbers at a Glance")

info_table(
    ["Item", "Number"],
    [
        ["Document classes", "12"],
        ["Visual format variants per class", "3-4 (37 total)"],
        ["Single-doc training PDFs", "12,000 (1,000 per class)"],
        ["Multi-page PDFs (continuation training)", "3,000 (2 pages each)"],
        ["Splitting packets", "5,000 (avg 5.3 pages, range 2-10)"],
        ["Total PDF files across all datasets", "20,000"],
        ["Total page images for training", "~44,500"],
        ["START labels", "~35,000 (78%)"],
        ["CONTINUATION labels", "~9,500 (22%)"],
        ["Train / Val / Test split", "80% / 10% / 10%"],
        ["Model", "Qwen2.5-VL-3B-Instruct (NOT 7B — 3B is sufficient for one-page-at-a-time task)"],
        ["Why 3B not 7B", "One page in = one label out. 3B fine-tuned beats 7B general. 2x faster inference."],
        ["Why Qwen2.5-VL not Gemma 3 4B", "DocVQA: 84% vs 69%. Qwen2.5-VL is document-specialist. Gemma is general VLM."],
        ["Fine-tuning method", "QLoRA (4-bit quantization + LoRA adapters)"],
        ["GPU needed", "1x RTX 3090/4090 (16-24GB) — runs on a single consumer GPU"],
        ["Estimated training time", "3-5 hours (3B trains faster than 7B)"],
        ["Output format", "<Document Class> | <START or CONTINUATION>"],
    ]
)

# Footer
doc.add_paragraph()
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run(
    f"DHL Document Intelligence Fine-Tuning Guide  |  {datetime.date.today().strftime('%d %B %Y')}\n"
    "Confidential — Internal Reference"
).font.size = Pt(9)

# Save
out = r"D:\finetuning\DHL_Document_finetuning\DHL_Finetuning_Complete_Guide.docx"
doc.save(out)
print(f"Saved: {out}")
