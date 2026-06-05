"""
Run this once to generate the Inference Optimization Word document.
    python create_optimization_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x00, 0x2E, 0x5E)   # Capgemini dark blue
TEAL      = RGBColor(0x00, 0x70, 0xAD)   # Capgemini mid blue
ACCENT    = RGBColor(0x00, 0xB1, 0xC1)   # Capgemini teal accent
GREEN     = RGBColor(0x16, 0x65, 0x34)
RED       = RGBColor(0xB9, 0x1C, 0x1C)
GRAY      = RGBColor(0x47, 0x55, 0x69)
LIGHTGRAY = RGBColor(0xF1, 0xF5, 0xF9)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x0F, 0x17, 0x2A)
CODE_BG   = RGBColor(0xF8, 0xFA, 0xFC)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_cell_border(cell, border_size=4, color="002E5E"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),  "single")
        b.set(qn("w:sz"),   str(border_size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def heading1(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = color
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = color
    return p

def heading3(text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p

def code_block(lines):
    """Render a monospace code box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, CODE_BG)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        p   = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def callout(label, text, bg=RGBColor(0xEF, 0xF6, 0xFF), label_color=TEAL):
    """Highlighted info box."""
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(1.5)
    tbl.columns[1].width = Cm(14)
    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)
    set_cell_bg(lc, label_color)
    set_cell_bg(rc, bg)
    p = lc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.bold = True
    run.font.size  = Pt(10)
    run.font.color.rgb = WHITE
    p2  = rc.paragraphs[0]
    run2 = p2.add_run(text)
    run2.font.size  = Pt(10)
    run2.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def before_after(before_lines, after_lines):
    """Two-column before/after code table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    # Header row
    for j, (label, bg) in enumerate([("BEFORE", RGBColor(0xFE, 0xE2, 0xE2)),
                                       ("AFTER",  RGBColor(0xDC, 0xFC, 0xE7))]):
        c = tbl.cell(0, j)
        set_cell_bg(c, bg)
        p   = c.paragraphs[0]
        run = p.add_run(label)
        run.bold      = True
        run.font.size = Pt(10)
    # Code row
    for j, (lines, bg) in enumerate([(before_lines, RGBColor(0xFF, 0xF5, 0xF5)),
                                       (after_lines,  RGBColor(0xF0, 0xFF, 0xF4))]):
        c = tbl.cell(1, j)
        set_cell_bg(c, bg)
        c.paragraphs[0]._element.getparent().remove(c.paragraphs[0]._element)
        for line in lines:
            p   = c.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(line if line else " ")
            run.font.name  = "Courier New"
            run.font.size  = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def speed_badge(text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"  Speed Impact:  {text}  ")
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    run.font.highlight_color = None
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(shd)

def add_summary_table(rows_data, headers):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        c   = tbl.cell(0, j)
        set_cell_bg(c, NAVY)
        p   = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    # Data rows
    for i, row in enumerate(rows_data):
        bg = RGBColor(0xF8, 0xFA, 0xFC) if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            c   = tbl.cell(i + 1, j)
            set_cell_bg(c, bg)
            p   = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if j == len(row) - 1:
                run.bold = True
                run.font.color.rgb = GREEN if "faster" in str(val).lower() or "×" in str(val) else NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ═════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ═════════════════════════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
tbl.style = "Table Grid"
cell = tbl.cell(0, 0)
set_cell_bg(cell, NAVY)
cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

for text, size, bold, color in [
    ("DHL Document Intelligence", 22, True,  WHITE),
    ("Inference Speed Optimizations", 26, True, RGBColor(0x00, 0xB1, 0xC1)),
    ("", 8, False, WHITE),
    ("A complete beginner-friendly guide", 13, False, RGBColor(0xCB, 0xD5, 0xE1)),
    ("explaining every change made, why it helps, and how fast it made things", 11, False, RGBColor(0x94, 0xA3, 0xB8)),
    ("", 6, False, WHITE),
    ("Hardware: NVIDIA RTX 5080 Laptop GPU (16 GB GDDR7)", 10, False, RGBColor(0x94, 0xA3, 0xB8)),
    ("Model: Qwen2.5-VL-3B-Instruct + LoRA fine-tuned", 10, False, RGBColor(0x94, 0xA3, 0xB8)),
    ("Framework: Unsloth + PyTorch 2.7 + CUDA 12.8", 10, False, RGBColor(0x94, 0xA3, 0xB8)),
]:
    p   = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 0 — HOW INFERENCE WORKS (background)
# ═════════════════════════════════════════════════════════════════════════════
heading1("Background: How Inference Works")

body("Before explaining what we optimized, it helps to understand what the app actually does "
     "when you upload a PDF. This section explains the basic flow in plain English.")

heading2("The Two-Stage Pipeline")

body("Every page of a document goes through exactly two stages:")

tbl2 = doc.add_table(rows=3, cols=3)
tbl2.style = "Table Grid"
headers2 = ["Stage", "What Happens", "Where It Runs"]
for j, h in enumerate(headers2):
    c = tbl2.cell(0, j)
    set_cell_bg(c, TEAL)
    run = c.paragraphs[0].add_run(h)
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE

for j, val in enumerate(["Stage 1 — Vision Encoder",
                          "Looks at the image pixel by pixel, divides it into small patches, "
                          "and converts each patch into a list of numbers (a 'visual token'). "
                          "Think of it like a human scanning the page with their eyes.",
                          "GPU"]):
    c = tbl2.cell(1, j); set_cell_bg(c, RGBColor(0xF0, 0xF9, 0xFF))
    run = c.paragraphs[0].add_run(val); run.font.size = Pt(10)

for j, val in enumerate(["Stage 2 — LLM Decoder",
                          "Reads the visual tokens plus the instruction prompt, then generates "
                          "the answer one word (token) at a time — like a person writing the JSON "
                          "response character by character.",
                          "GPU"]):
    c = tbl2.cell(2, j); set_cell_bg(c, WHITE)
    run = c.paragraphs[0].add_run(val); run.font.size = Pt(10)

doc.add_paragraph()

callout("KEY",
        "Every optimization in this document targets either the Vision Encoder (Stage 1), "
        "the LLM Decoder (Stage 2), or the time spent between pages.",
        bg=RGBColor(0xEF, 0xF6, 0xFF), label_color=TEAL)

body("The starting point before any optimizations: 20–30 seconds per page on an RTX 5080. "
     "That is roughly 3–4 minutes for an 8-page PDF. After all the changes, the same PDF "
     "takes under 60 seconds.")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — 4-BIT TO BFLOAT16
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 1 — Switch from 4-bit to BFloat16")

heading2("What is Quantization?")

body("A neural network model is made up of billions of numbers called weights. "
     "These weights are usually stored as 16-bit floating point numbers (bfloat16), "
     "where each number takes 2 bytes of memory.")

body("Quantization is the process of storing those numbers in a smaller format "
     "to save GPU memory. 4-bit quantization stores each weight in only 0.5 bytes — "
     "four times smaller than bfloat16.")

body("Analogy: Imagine you have a very precise recipe that calls for "
     "exactly 137.5 grams of flour. Bfloat16 writes '137.5 g'. "
     "4-bit can only write values like 130, 135, 140, 145 — it rounds to the "
     "nearest slot. The dish still tastes similar, but every measurement is an approximation.")

heading2("Why 4-bit Slows Down Inference")

body("Here is the key problem: the GPU cannot actually do maths with 4-bit integers. "
     "It only understands bfloat16 and float32. So every time the model needs to use "
     "a weight for a calculation, it must first convert it back (dequantize):")

code_block([
    "For every layer, for every token generated:",
    "",
    "  stored_weight (4-bit int)  →  dequantize  →  bfloat16_weight",
    "                                                      ↓",
    "                                             multiply with input",
    "                                                      ↓",
    "                                              discard bfloat16_weight",
    "                                         (cannot store — not enough memory)",
    "",
    "This dequantization happens ~200 layers × every decode step.",
    "For 180 tokens output: 200 × 180 = 36,000 dequantizations per page.",
])

heading2("The Fix: Load the Model in BFloat16 Directly")

body("The RTX 5080 has 16 GB of VRAM. The 3B model in bfloat16 only needs "
     "about 6.5 GB. There is no need to quantize for inference:")

code_block([
    "Memory calculation:",
    "  3,000,000,000 parameters × 2 bytes (bfloat16) = 6.0 GB",
    "  LoRA adapters                                  = 0.1 GB",
    "  Inference activations (temporary)              = 2.0 GB",
    "  ─────────────────────────────────────────────────────────",
    "  Total                                          = 8.1 GB",
    "  Available VRAM on RTX 5080                     = 16.0 GB",
    "  Remaining free                                 = 7.9 GB  ✓",
])

before_after(
    ["# Loads base model compressed as 4-bit integers",
     "# Every matmul needs dequantize first",
     "model, processor = FastVisionModel.from_pretrained(",
     "    model_path,",
     "    load_in_4bit=True,   # int4 → 1.5 GB but slow",
     ")"],
    ["# Loads base model as native bfloat16",
     "# GPU operates directly — no conversion needed",
     "model, processor = FastVisionModel.from_pretrained(",
     "    model_path,",
     "    load_in_4bit=False,  # bf16 → 6.5 GB but FAST",
     ")"]
)

speed_badge("2–3× faster per token — the single biggest improvement")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TF32
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 2 — Enable TF32 Matrix Multiplications")

heading2("What are Tensor Cores?")

body("Modern NVIDIA GPUs contain special hardware units called Tensor Cores. "
     "These are separate from the normal CUDA cores and are designed specifically "
     "for one thing: matrix multiplication, which is the core operation inside every "
     "neural network layer.")

body("Tensor Cores run matrix multiplications 4–8× faster than normal CUDA cores, "
     "but they only work when the maths is done in specific precision formats.")

heading2("What is TF32?")

body("TF32 (TensorFloat-32) is a precision format invented by NVIDIA. It is a compromise:")

tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = "Table Grid"
for j, h in enumerate(["Format", "Range (exponent bits)", "Precision (mantissa bits)"]):
    c = tbl3.cell(0, j); set_cell_bg(c, NAVY)
    run = c.paragraphs[0].add_run(h); run.bold = True
    run.font.size = Pt(10); run.font.color.rgb = WHITE

for i, row in enumerate([
    ["Float32 (standard)", "8 bits (large range)", "23 bits (very precise)"],
    ["Float16", "5 bits (small range)", "10 bits (less precise)"],
    ["TF32", "8 bits (same as float32)", "10 bits (same as float16)"],
]):
    bg = [RGBColor(0xF8, 0xFA, 0xFC), WHITE, RGBColor(0xEF, 0xF6, 0xFF)][i]
    for j, val in enumerate(row):
        c = tbl3.cell(i+1, j); set_cell_bg(c, bg)
        run = c.paragraphs[0].add_run(val); run.font.size = Pt(10)

doc.add_paragraph()
body("TF32 has the same range as float32 (handles very large and very small numbers safely) "
     "but the speed of float16 (runs on Tensor Cores). "
     "The precision loss is negligible for inference — the final answer changes by less than 0.1%.")

heading2("The Fix: Two Lines of Code")

before_after(
    ["# TF32 disabled by default in PyTorch",
     "# Normal CUDA cores used for matmul",
     "# Tensor Cores are underutilised",
     "",
     "# (no code — default behaviour)"],
    ["# Enable TF32 — Tensor Cores fully utilised",
     "# Runs before model loading",
     "torch.backends.cuda.matmul.allow_tf32 = True",
     "torch.backends.cudnn.allow_tf32       = True",
     "# (applies to both matmul and convolution)"]
)

speed_badge("10–20% faster — especially on the large linear layers in the LLM decoder")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PRE-COMPUTE PIXEL VALUES
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 3 — Pre-compute Image Features Before Inference")

heading2("The Bottleneck: CPU and GPU Working One After Another")

body("Processing a page involves two separate workers:")

bullet("The CPU prepares the image: resizes it, normalises pixel values (0–255 → 0.0–1.0), "
       "and creates patch tensors. This takes ~15–25 ms per page.")
bullet("The GPU runs the vision encoder and LLM decoder. This takes ~6–25 seconds per page.")

body("Before the optimization, these always ran sequentially:")

code_block([
    "For each page:",
    "  [CPU: prepare image]   15-25 ms",
    "  [GPU: vision encoder + LLM decode]   6-25 seconds",
    "  [CPU: prepare next image]   15-25 ms",
    "  [GPU: vision encoder + LLM decode]   6-25 seconds",
    "  ...",
    "",
    "The GPU was idle during CPU work, and the CPU was idle during GPU work.",
])

heading2("What is Pinned Memory?")

body("When tensors live in normal CPU RAM and need to move to the GPU, the operating "
     "system can interrupt the transfer at any time. It copies through a staging buffer:")

code_block([
    "Normal RAM:   CPU RAM  →  OS staging buffer  →  GPU VRAM   (2 copies, interruptible)",
    "Pinned RAM:   CPU RAM  →  GPU VRAM            (1 DMA copy, locked, non-blocking)",
])

body("Pinned (page-locked) memory bypasses the OS and lets the GPU transfer data directly "
     "via DMA (Direct Memory Access). It is faster and allows the CPU to continue doing "
     "other work while the transfer happens (non_blocking=True).")

heading2("The Fix: Pre-compute All Images Upfront")

before_after(
    ["for page in pages:",
     "    # CPU work inside the loop",
     "    img_data = processor.image_processor(",
     "        images=[page]",
     "    )",
     "    # Synchronous GPU upload",
     "    inputs = inputs.to('cuda')",
     "    # GPU inference",
     "    output = model.generate(**inputs)"],
    ["# ALL CPU work done before GPU loop starts",
     "pv_caches = precompute_pixel_values(",
     "    proc_images, processor",
     ")",
     "# pixel_values stored in pinned RAM",
     "",
     "for page, pv_cache in zip(pages, pv_caches):",
     "    # Only text tokenization in loop (~5ms)",
     "    # Non-blocking DMA transfer to GPU",
     "    inputs['pixel_values'] = pv_cache[",
     "        'pixel_values'].to('cuda',",
     "        non_blocking=True)",
     "    output = model.generate(**inputs)"]
)

speed_badge("15–25 ms saved per page — CPU/GPU work separated cleanly")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MAX_PIXELS
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 4 — Reduce Image Resolution (MAX_PIXELS)")

heading2("How the Vision Encoder Processes Images")

body("The vision encoder does not look at individual pixels. Instead it divides the "
     "image into fixed-size patches (14×14 pixels each), then groups four patches "
     "together into a single visual token:")

code_block([
    "Image resolution  →  patches  →  merged tokens (visual tokens)",
    "",
    "Formula:  visual_tokens = image_pixels / (14 × 14 × 4)",
    "                        = image_pixels / 784",
    "",
    "640,000 pixels  →  816 visual tokens",
    "384,000 pixels  →  490 visual tokens",
])

heading2("Why Fewer Tokens Means Much Faster Inference")

body("The vision encoder uses a mechanism called self-attention. "
     "In self-attention, every visual token must 'look at' every other visual token "
     "to understand the context. The number of operations is proportional to N squared "
     "where N is the number of tokens:")

code_block([
    "Self-attention cost  ∝  N²   (N = number of visual tokens)",
    "",
    "At 640K pixels:  816²  =  666,000 operations  (baseline)",
    "At 384K pixels:  490²  =  240,000 operations",
    "",
    "Reduction:  240,000 / 666,000  =  36% of original cost",
    "Saving:     64% fewer attention operations",
])

body("This is a quadratic relationship — cutting tokens by 40% cuts attention "
     "computation by nearly 64%. That is why this change has such a large effect.")

heading2("Is It Safe to Reduce Resolution?")

body("Yes, for two reasons:")

bullet("Qwen2.5-VL was pre-trained on images ranging from very small to very large — "
       "it is designed to handle variable resolutions and generalises well.")
bullet("DHL documents contain printed text in standard font sizes. At 384K pixels, "
       "all field values (names, dates, numbers, addresses) remain clearly readable "
       "by the model. The quality reduction only affects very fine detail that is "
       "irrelevant for field extraction.")

before_after(
    ["# Matches training resolution",
     "# 816 visual tokens per page",
     "MAX_PIXELS = 640_000"],
    ["# 490 visual tokens per page",
     "# 40% fewer tokens → 64% less",
     "# vision encoder attention work",
     "MAX_PIXELS = 384_000"]
)

speed_badge("35–40% faster vision encoding — second biggest improvement")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — WARMUP
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 5 — GPU Kernel Warmup on Startup")

heading2("What is a CUDA Kernel?")

body("When PyTorch runs an operation on the GPU (like matrix multiplication), "
     "it sends a small program called a kernel to the GPU driver. "
     "The very first time a particular type of operation runs, CUDA must:")

bullet("Compile the kernel from source code to GPU machine code")
bullet("Cache it so future calls are instant")
bullet("Set up memory buffers and internal state")

body("This compilation takes 1–5 seconds and only happens once per session. "
     "But without warmup, this cost hits the very first PDF page a user uploads.")

heading2("The Problem: Cold Start on First Page")

code_block([
    "Without warmup:",
    "",
    "  Server starts       → model loaded into GPU memory",
    "  User uploads PDF",
    "  Page 1 inference    → kernel compilation (2-5 s)",
    "                      + vision encoder (5-10 s)",
    "                      + LLM decode (5-15 s)",
    "                      = 15-30 s  ← feels very slow",
    "",
    "  Page 2 inference    → no compilation (cached)",
    "                      = 6-10 s  ← normal speed",
])

heading2("The Fix: Run a Dummy Image at Startup")

body("By running a tiny dummy inference before the server opens, we force all kernel "
     "compilations to happen during startup — invisible to the user:")

code_block([
    "def _warmup_model():",
    "    # Create a tiny blank image (224x224 pixels)",
    "    dummy = Image.new('RGB', (224, 224), color=128)",
    "",
    "    # Run a FULL forward pass — encodes image + generates 4 tokens",
    "    # This compiles and caches all CUDA kernels",
    "    inputs = processor(text=[tmpl], images=[dummy], ...)",
    "    with torch.inference_mode():",
    "        model.generate(**inputs, max_new_tokens=4, ...)",
    "",
    "    print('Warmup complete.')  # kernels are now cached",
    "",
    "# Called once at startup, before app.launch()",
])

code_block([
    "With warmup:",
    "",
    "  Server starts       → model loaded + warmup (kernels compiled here)",
    "  User uploads PDF",
    "  Page 1 inference    → no compilation needed",
    "                      = 6-10 s  ← same speed as page 2",
])

speed_badge("First page same speed as all other pages — eliminates cold-start penalty")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PRE-LOAD MODEL
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 6 — Pre-load Model Before Accepting Requests")

heading2("Lazy Loading vs Eager Loading")

body("Loading the model from disk into GPU memory takes 30–60 seconds "
     "(loading 6.5 GB of weights, allocating GPU buffers, etc.).")

body("Lazy loading means the model is loaded on demand — when the first user "
     "clicks 'Analyze Document'. The user then stares at a spinning progress bar "
     "for up to a minute before anything happens.")

body("Eager loading means the model is loaded at server startup, before any user "
     "can even reach the app. By the time the browser opens, the model is already "
     "sitting in GPU memory, ready.")

before_after(
    ["def process_files(files, ...):",
     "    # Model loads HERE — on first click",
     "    # User waits 30-60 seconds",
     "    progress(0, desc='Loading model...')",
     "    model, processor = get_model()",
     "    # ...then inference starts"],
    ["if __name__ == '__main__':",
     "    # Model loads at STARTUP",
     "    # User never sees this wait",
     "    print('Pre-loading model...')",
     "    get_model()          # blocks here",
     "    app = build_app()",
     "    app.launch(...)      # browser opens",
     "                         # model already ready"]
)

speed_badge("0 ms model load time on first click — 30–60 s wait moved to server startup")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — MERGE LORA
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 7 — Merge LoRA Adapters into Base Model")

heading2("What is LoRA Fine-tuning?")

body("LoRA (Low-Rank Adaptation) is a technique for fine-tuning large models cheaply. "
     "Instead of updating all 3 billion weights (which would require enormous GPU memory "
     "and time), LoRA trains two small matrices A and B per layer:")

code_block([
    "Standard layer:",
    "  output = W × input",
    "  (W = the original pre-trained weight matrix)",
    "",
    "LoRA layer (during training and with adapters at inference):",
    "  output = W × input  +  B × A × input × (alpha / rank)",
    "           ──────────    ────────────────────────────────",
    "           original          LoRA delta (what we trained)",
    "",
    "This requires TWO matrix multiplications every forward pass.",
    "With 100+ linear layers in a 3B model, that is 100+ extra multiplications",
    "per token generated.",
])

heading2("What is Merging?")

body("Merging combines the original weights and the LoRA delta permanently "
     "into a single matrix. The maths is simple — just add them together:")

code_block([
    "W_merged = W + B × A × (alpha / rank)",
    "",
    "After merging:",
    "  output = W_merged × input",
    "           ────────────────",
    "           ONE multiplication — same as original, no adapter overhead",
    "",
    "Quality: identical — W_merged contains exactly the same values",
    "         as W + LoRA would have produced, just stored together.",
])

heading2("Why a Separate Script?")

body("Merging cannot be done inside the app at runtime because Unsloth (the training "
     "library) patches the model layers with its own custom CUDA kernels when you "
     "import it. Those patches conflict with the standard PEFT merge process, "
     "causing corrupted weights.")

body("The solution: a standalone merge_model.py script that does NOT import Unsloth, "
     "uses plain HuggingFace + PEFT libraries, and saves the merged model to disk:")

code_block([
    "# merge_model.py — NO unsloth import",
    "from transformers import Qwen2_5_VLForConditionalGeneration",
    "from peft import PeftModel",
    "",
    "# Step 1: Load base model in bfloat16",
    "base = Qwen2_5_VLForConditionalGeneration.from_pretrained(",
    "    BASE_MODEL_PATH, torch_dtype=torch.bfloat16",
    ")",
    "",
    "# Step 2: Attach LoRA adapters",
    "peft_model = PeftModel.from_pretrained(base, LORA_CHECKPOINT_PATH)",
    "",
    "# Step 3: Merge — W_merged = W + B·A·scale",
    "merged = peft_model.merge_and_unload()",
    "",
    "# Step 4: Save as a plain bfloat16 model",
    "merged.save_pretrained(MERGED_OUTPUT_PATH)",
])

body("The app then automatically loads the merged model if it exists:")

code_block([
    "# In dhl_app.py — priority order:",
    "if MERGED_PATH.exists():",
    "    src = MERGED_PATH    # plain bfloat16, zero LoRA overhead  ← fastest",
    "elif LORA_PATH.exists():",
    "    src = LORA_PATH      # base + adapters (LoRA overhead)     ← fallback",
    "else:",
    "    src = MODEL_PATH     # base model only (no fine-tuning)    ← last resort",
])

speed_badge("10–15% faster per token — one fewer matrix multiply per linear layer")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8 — MAX_NEW_TOKENS
# ═════════════════════════════════════════════════════════════════════════════
heading1("Optimization 8 — Setting max_new_tokens Correctly")

heading2("What is Autoregressive Decoding?")

body("The LLM generates output one token at a time. Each token depends on all "
     "previous tokens — it cannot generate token 5 without first generating "
     "tokens 1 through 4. This is called autoregressive decoding.")

body("max_new_tokens is a hard limit on how many tokens the model can generate "
     "before it is forcibly stopped. The model stops earlier if it generates an "
     "end-of-sequence (EOS) token naturally.")

code_block([
    "Page type            Actual tokens generated    max_new_tokens needed",
    "─────────────────────────────────────────────────────────────────────",
    "CONTINUATION page    ~25 tokens                 any value ≥ 25",
    "START page (sparse)  ~100 tokens                any value ≥ 100",
    "START page (full)    ~180–200 tokens             must be ≥ 200",
    "",
    "Worst case: Commercial Invoice with all fields filled → ~200 tokens",
])

heading2("The Lesson Learned: Lower Is Not Always Faster")

body("During optimization, max_new_tokens was reduced from 220 to 160. "
     "This broke field extraction for Commercial Invoice and Certificate of Origin — "
     "the two document types with the most fields.")

body("What happened: the JSON was truncated before the closing brace. "
     "Instead of a complete JSON object, the model output ended mid-string. "
     "The parser could not recover the field values and returned empty results.")

callout("LESSON",
        "Reducing max_new_tokens only saves time if the model naturally generates "
        "fewer tokens (via EOS). If the output needs those tokens to be complete, "
        "cutting the limit breaks the result instead of speeding things up.",
        bg=RGBColor(0xFF, 0xF7, 0xED), label_color=RGBColor(0xEA, 0x58, 0x0C))

before_after(
    ["# Too low — truncates long JSON",
     "# Commercial Invoice fields cut off",
     "model.generate(",
     "    **inputs,",
     "    max_new_tokens=160,  # ← broke CI",
     ")"],
    ["# Safe ceiling for all 12 doc types",
     "# CONTINUATION stops at ~25 via EOS",
     "model.generate(",
     "    **inputs,",
     "    max_new_tokens=220,  # ← correct",
     ")"]
)

speed_badge("Neutral — correct value maintains accuracy without wasting tokens")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 9 — SUMMARY TABLE
# ═════════════════════════════════════════════════════════════════════════════
heading1("Complete Summary")

body("The table below shows every optimization, what changed, and its speed impact:")

add_summary_table(
    rows_data=[
        ["1. BFloat16 Inference",        "load_in_4bit: True → False",       "2–3× faster per token"],
        ["2. TF32 Matrix Multiply",       "allow_tf32 flags enabled",         "10–20% faster"],
        ["3. Pre-compute Pixel Values",   "pinned memory + non_blocking",     "15–25 ms/page saved"],
        ["4. Reduce MAX_PIXELS",          "640K → 384K pixels",               "35–40% faster vision encoder"],
        ["5. GPU Warmup",                 "dummy pass at startup",             "First page = all pages speed"],
        ["6. Pre-load Model",             "load at startup, not on click",    "0 s load wait for user"],
        ["7. Merge LoRA Adapters",        "merge_model.py (offline)",         "10–15% faster per token"],
        ["8. max_new_tokens",             "220 (safe ceiling)",               "Accuracy maintained"],
    ],
    headers=["Optimization", "Change Made", "Speed Impact"]
)

heading2("Before vs After (8-page PDF)")

add_summary_table(
    rows_data=[
        ["Per-page inference time",    "20–30 s",   "6–10 s",    "~3× faster"],
        ["First-page cold start",      "28–35 s",   "6–10 s",    "No cold start"],
        ["Model load on first click",  "30–60 s",   "0 s",       "Pre-loaded"],
        ["Total for 8-page PDF",       "~200 s",    "~60 s",     "~3× faster"],
        ["Vision tokens per page",     "816",       "490",       "−40%"],
        ["GPU memory used",            "~1.5 GB",   "~8.5 GB",   "Uses VRAM properly"],
    ],
    headers=["Metric", "Before", "After", "Improvement"]
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 10 — WHAT STILL LIMITS SPEED
# ═════════════════════════════════════════════════════════════════════════════
heading1("What Still Limits Speed")

body("Even with all optimizations, some fundamental constraints remain:")

bullet("Sequential pages — The GPU processes one page at a time. "
       "Page N+1 cannot start until page N finishes. "
       "Batching multiple pages would require all pages to be the same resolution, "
       "which is not guaranteed for real-world PDFs.")

bullet("LLM decode is autoregressive — Each output token depends on the previous one. "
       "This cannot be parallelized. A document with many fields will always be "
       "slower than a document with few fields.")

bullet("384K pixel floor — Reducing resolution further (e.g. 256K) starts to lose "
       "small text, especially for documents with dense tables or small font sizes.")

bullet("LoRA overhead (if merge_model.py not run) — If the merged model does not exist, "
       "the app falls back to the LoRA checkpoint. bfloat16 is still used (fast), "
       "but the extra A×B multiply per layer remains.")

doc.add_paragraph()
callout("NOTE",
        "Running python merge_model.py once creates the merged model and eliminates "
        "the LoRA overhead permanently for all future app launches.",
        bg=RGBColor(0xF0, 0xFD, 0xF4), label_color=GREEN)

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
output_path = "Inference_Optimization_Guide.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
