"""Extract full text from both existing papers."""
from docx import Document
import os

def extract(path, label):
    if not os.path.exists(path):
        print(f"\n{'='*60}\n{label}: FILE NOT FOUND\n{'='*60}")
        return
    doc = Document(path)
    print(f"\n{'='*60}\n{label}\n{'='*60}")
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"  [{para.style.name}]  {para.text}")
    for i, tbl in enumerate(doc.tables):
        print(f"\n  [TABLE {i}]")
        for row in tbl.rows:
            cells = [c.text.strip()[:50] for c in row.cells]
            print("  | " + " | ".join(cells))

extract(r"D:\finetuning\DHL_Document_finetuning\Technical_Paper_Document_Intelligence.docx",
        "ORIGINAL PAPER")
extract(r"D:\finetuning\DHL_Document_finetuning\Technical_Paper_Document_Intelligence_v3.docx",
        "V3 PAPER")
