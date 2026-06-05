"""Extract full text from existing technical paper."""
from docx import Document
doc = Document(r"D:\finetuning\DHL_Document_finetuning\Technical_Paper_Document_Intelligence.docx")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{para.style.name}] {para.text}")
for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLE {ti} ---")
    for row in table.rows:
        print(" | ".join(cell.text[:60] for cell in row.cells))
