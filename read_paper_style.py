"""Extract text content and color theme from existing Technical_Paper_Document_Intelligence.docx"""
from docx import Document
from docx.shared import RGBColor
import json

doc = Document(r"D:\finetuning\DHL_Document_finetuning\Technical_Paper_Document_Intelligence.docx")

# Extract all paragraph styles and colors
colors_found = set()
print("=== PARAGRAPHS (first 300) ===")
for i, para in enumerate(doc.paragraphs[:300]):
    if para.text.strip():
        style_name = para.style.name
        # Get run-level colors
        for run in para.runs:
            if run.font.color and run.font.color.type:
                try:
                    rgb = run.font.color.rgb
                    colors_found.add(str(rgb))
                except:
                    pass
        # Get paragraph background
        print(f"[{i:03d}] [{style_name}] {para.text[:120]}")

print("\n=== TABLE CONTENT ===")
for ti, table in enumerate(doc.tables[:5]):
    print(f"\nTable {ti}:")
    for row in table.rows[:3]:
        for cell in row.cells[:4]:
            print(f"  | {cell.text[:40]:<40}", end="")
        print()

print("\n=== COLORS FOUND IN RUNS ===")
for c in sorted(colors_found):
    print(f"  #{c}")

# Check theme colors in document XML
print("\n=== DOCUMENT XML THEME SNIPPET ===")
try:
    theme_xml = doc.part.theme_part._element.xml[:2000]
    print(theme_xml)
except:
    print("No theme part found")
